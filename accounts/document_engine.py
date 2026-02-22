import io
from docxtpl import DocxTemplate
from django.utils import timezone
from .models import DocumentTemplate, Student, Order
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.db.models import Count, Q


class DocumentGenerator:
    @staticmethod
    def generate_document(template_id, object_id):
        template_obj = DocumentTemplate.objects.get(id=template_id)
        doc = DocxTemplate(template_obj.file.path)
        
        context = {}
        filename = "document.docx"

        if template_obj.context_type == 'STUDENT_CERT':
            student = Student.objects.get(id=object_id)
            context = {
                'date_today': timezone.now().strftime('%d.%m.%Y'),
                'full_name': student.user.get_full_name(),
                'first_name': student.user.first_name,
                'last_name': student.user.last_name,
                'student_id': student.student_id,
                'course': student.course,
                'group_name': student.group.name if student.group else 'Не назначена',
                'specialty_code': student.specialty.code if student.specialty else '',
                'specialty_name': student.specialty.name if student.specialty else '',
                'financing': student.get_financing_type_display(),
                'education_type': student.get_education_type_display(),
                'faculty_name': student.specialty.department.faculty.name if student.specialty else '',
                'institute_name': student.specialty.department.faculty.institute.name if student.specialty else '',
            }
            filename = f"Spravka_{student.user.last_name}.docx"

        elif template_obj.context_type == 'STUDENT_ORDER':
            order = StudentOrder.objects.get(id=object_id)
            student = order.student
            context = {
                'order_number': order.number,
                'order_date': order.date.strftime('%d.%m.%Y'),
                'order_reason': order.reason,
                'order_type': order.get_order_type_display(),
                'student_full_name': student.user.get_full_name(),
                'group_name': student.group.name if student.group else '',
                'course': student.course,
                'specialty_name': student.specialty.name if student.specialty else '',
                'initiator_name': order.initiated_by.get_full_name() if order.initiated_by else '',
            }
            filename = f"Prikaz_{order.number}_{student.user.last_name}.docx"

        doc.render(context)
        
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return file_stream, filename
    @staticmethod
    def generate_contingent_report(faculty=None):
        from .models import Group, Student
        
        doc = Document()
        
        section = doc.sections[0]
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

        
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("ОТЧЕТ О ДВИЖЕНИИ КОНТИНГЕНТА СТУДЕНТОВ\n")
        run.bold = True
        run.font.size = Pt(14)
        
        if faculty:
            title.add_run(f"Факультет: {faculty.name}\n").bold = True
        
        title.add_run(f"Дата формирования: {timezone.now().strftime('%d.%m.%Y')}")

        doc.add_paragraph()

        groups = Group.objects.all()
        if faculty:
            groups = groups.filter(specialty__department__faculty=faculty)

        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Курс'
        hdr_cells[1].text = 'Группа'
        hdr_cells[2].text = 'Активные'
        hdr_cells[3].text = 'В академ. отпуске'
        hdr_cells[4].text = 'Отчислены'
        hdr_cells[5].text = 'Выпускники'

        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        total_active = total_leave = total_expelled = total_grad = 0

        for group in groups.order_by('course', 'name'):
            students = Student.objects.filter(group=group)
            active = students.filter(status='ACTIVE').count()
            leave = students.filter(status='ACADEMIC_LEAVE').count()
            
            from .models import OrderItem
            expelled = OrderItem.objects.filter(student__status='EXPELLED', order__order_type='EXPEL', student__specialty=group.specialty).count()
            graduated = OrderItem.objects.filter(student__status='GRADUATED', order__order_type='GRADUATE', student__specialty=group.specialty).count()

            row_cells = table.add_row().cells
            row_cells[0].text = str(group.course)
            row_cells[1].text = group.name
            row_cells[2].text = str(active)
            row_cells[3].text = str(leave)
            row_cells[4].text = str(expelled)
            row_cells[5].text = str(graduated)

            total_active += active
            total_leave += leave
            total_expelled += expelled
            total_grad += graduated

        footer_cells = table.add_row().cells
        footer_cells[0].merge(footer_cells[1])
        footer_cells[0].text = 'ИТОГО:'
        footer_cells[0].paragraphs[0].runs[0].bold = True
        footer_cells[2].text = str(total_active)
        footer_cells[3].text = str(total_leave)
        footer_cells[4].text = str(total_expelled)
        footer_cells[5].text = str(total_grad)

        for cell in footer_cells:
            if cell.text:
                cell.paragraphs[0].runs[0].bold = True

        doc.add_paragraph("\nПодпись ответственного лица: ___________________")

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        filename = f"Contingent_Report_{timezone.now().strftime('%Y%m%d')}.docx"
        return file_stream, filename
