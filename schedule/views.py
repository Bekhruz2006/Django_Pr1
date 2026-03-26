from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib import messages
from django.http import HttpResponse, JsonResponse
from django.db.models import Q
from django.views.decorators.http import require_POST
from django.db import transaction
from django.utils import translation
from datetime import datetime, timedelta, date
import json
import uuid
import re
from io import BytesIO 
from django.db.models import Q, Sum
from django.urls import reverse
import logging
logger = logging.getLogger(__name__)
from .models import UnusedHourPool
from schedule.models import UnusedHourPool
from django.utils import timezone
logger = logging.getLogger('schedule')


try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
from schedule.models import ROOM_TYPES
from lms.models import Assignment

from django import forms
from .models import Subject, CreditType, CreditTemplate, ScheduleSlot, Semester, Classroom, TimeSlot, TeacherUnavailableSlot, AcademicPlan, PlanDiscipline, SubjectTemplate, SubjectMaterial, Building, Institute
from .forms import SubjectForm, RupImportForm, SemesterForm, ClassroomForm, BulkClassroomForm, TimeSlotGeneratorForm, MaterialUploadForm, ScheduleImportForm, AcademicPlanForm, PlanDisciplineForm, SubjectTemplateForm, BuildingForm, CreditTemplateForm
from .services import ScheduleImporter, RupImporter
from accounts.models import Group, Student, Teacher, Director, ProRector, Department, Faculty
import math
from django.utils.translation import gettext as _
from schedule.models import ROOM_TYPES
from .services import AIAssignmentService, AlgorithmicAssignmentService 
from .timetable_bridge import AutoScheduleEngineCpp as AutoScheduleEngine
from schedule.models import ScheduleException as SchedExc
from datetime import date as date_cls, datetime as dt_cls


import logging
logger = logging.getLogger(__name__)

def is_dean(user):
    return user.is_authenticated and hasattr(user, 'dean_profile')

def is_teacher(user):
    return user.is_authenticated and hasattr(user, 'teacher_profile')

def is_student(user):
    return user.is_authenticated and hasattr(user, 'student_profile')

def get_time_slots_for_shift(shift, institute=None):
    qs = TimeSlot.objects.filter(shift=shift)
    if institute:
        inst_qs = qs.filter(institute=institute)
        if inst_qs.exists():
            return inst_qs.order_by('start_time')
    return qs.filter(institute__isnull=True).order_by('start_time')


def is_dean_or_admin(user):
    return user.is_authenticated and (
        user.is_superuser or 
        user.role in ['RECTOR', 'PRO_RECTOR', 'DIRECTOR', 'DEAN', 'VICE_DEAN', 'HEAD_OF_DEPT']
    )

def institute_pair_ratio(institute):
    if not institute:
        return 2.0
    acad = getattr(institute, 'academic_hour_duration', None) or 50
    pair = getattr(institute, 'pair_duration', None) or 50
    try:
        return float(pair) / float(acad)
    except (TypeError, ValueError, ZeroDivisionError):
        return 2.0

def is_facility_admin(user):
    return user.is_authenticated and (
        user.is_superuser or
        user.role in ['RECTOR', 'PRO_RECTOR', 'DIRECTOR', 'DEAN', 'VICE_DEAN']
    )

def is_dept_head_or_above(user):
    return user.is_authenticated and (
        user.is_superuser or 
        user.role in['RECTOR', 'PRO_RECTOR', 'DIRECTOR', 'DEAN', 'VICE_DEAN', 'HEAD_OF_DEPT']
    )

def safe_int(val) -> int:
    if val is None:
        return 0
    if isinstance(val, bool):
        return int(val)
    if isinstance(val, (int, float)):
        return int(val)
    try:
        s = str(val).strip()
        if s.lower() in ('', 'none', 'null', '-', '—', 'нет', 'н/д', 'n/a'):
            return 0
        try:
            return int(float(s.replace(',', '.')))
        except ValueError:
            pass
        m = re.search(r'(\d+(?:[.,]\d+)?)', s)
        if m:
            return int(float(m.group(1).replace(',', '.')))
        return 0
    except (ValueError, TypeError, AttributeError):
        return 0

def get_active_semester_for_group(group):
    if group and group.specialty and group.specialty.department.faculty:
        faculty = group.specialty.department.faculty
        semester = Semester.objects.filter(faculty=faculty, is_active=True, course=group.course).first()
        if semester:
            return semester
    semester = Semester.objects.filter(course=group.course, is_active=True).first()
    if not semester:
        semester = Semester.objects.filter(is_active=True).first()
        
    return semester

@login_required
def schedule_constructor(request):
    if not is_dean_or_admin(request.user):
        messages.error(request, _("Доступ запрещен"))
        return redirect('core:dashboard')

    groups = Group.objects.all().select_related('specialty').order_by('course', 'name')

    if hasattr(request.user, 'dean_profile') or hasattr(request.user, 'vicedean_profile'):
        profile = getattr(request.user, 'dean_profile', None) or getattr(request.user, 'vicedean_profile', None)
        faculty = profile.faculty
        groups = groups.filter(specialty__department__faculty=faculty)
    elif hasattr(request.user, 'director_profile') or hasattr(request.user, 'prorector_profile'):
        profile = getattr(request.user, 'director_profile', None) or getattr(request.user, 'prorector_profile', None)
        institute = profile.institute
        if institute:
            groups = groups.filter(specialty__department__faculty__institute=institute)

    if 'group' in request.GET:
        selected_group_id = request.GET.get('group')
        if selected_group_id:
            request.session['last_constructor_group'] = selected_group_id
        else:
            request.session.pop('last_constructor_group', None)
    else:
        selected_group_id = request.session.get('last_constructor_group')

    selected_group = None
    if selected_group_id:
        selected_group = Group.objects.filter(id=selected_group_id).first()
        if not selected_group:
            request.session.pop('last_constructor_group', None)

    if 'semester' in request.GET:
        semester_id = request.GET.get('semester')
        if semester_id:
            request.session['last_constructor_semester'] = semester_id
        else:
            request.session.pop('last_constructor_semester', None)
    else:
        semester_id = request.session.get('last_constructor_semester')

    active_semester = None
    if semester_id:
        active_semester = Semester.objects.filter(id=semester_id).first()
        if not active_semester:
            request.session.pop('last_constructor_semester', None)
            
    if not active_semester and selected_group:
        active_semester = get_active_semester_for_group(selected_group)
    else:
        if hasattr(request.user, 'dean_profile') or hasattr(request.user, 'vicedean_profile'):
            profile = getattr(request.user, 'dean_profile', None) or getattr(request.user, 'vicedean_profile', None)
            faculty = profile.faculty
            active_semester = Semester.objects.filter(Q(faculty=faculty) | Q(faculty__isnull=True), is_active=True).first()
        elif hasattr(request.user, 'director_profile') or hasattr(request.user, 'prorector_profile'):
            profile = getattr(request.user, 'director_profile', None) or getattr(request.user, 'prorector_profile', None)
            institute = profile.institute
            if institute:
                active_semester = Semester.objects.filter(Q(faculty__institute=institute) | Q(faculty__isnull=True), is_active=True).first()
        
        if not active_semester:
            active_semester = Semester.objects.filter(is_active=True).first()

    if not active_semester:
        active_semester = Semester.objects.order_by('-start_date').first()
        if not active_semester:
            messages.warning(request, _('Сначала создайте семестр в настройках.'))
            return redirect('schedule:manage_semesters')

    schedule_data = {}
    subjects_to_schedule = []

    active_plan = None
    if selected_group:
        active_plan = AcademicPlan.objects.filter(group=selected_group, is_active=True).order_by('-admission_year').first()
        if not active_plan and selected_group.specialty:
            active_plan = AcademicPlan.objects.filter(specialty=selected_group.specialty, is_active=True).order_by('-admission_year').first()

        slots_qs = ScheduleSlot.objects.filter(
            group=selected_group,
            semester=active_semester,
            is_active=True
        ).select_related('subject', 'teacher__user', 'time_slot', 'classroom')
        slots = list(slots_qs)

        for slot in slots:
            if slot.day_of_week not in schedule_data:
                schedule_data[slot.day_of_week] = {}
            if slot.time_slot.id not in schedule_data[slot.day_of_week]:
                schedule_data[slot.day_of_week][slot.time_slot.id] = []
            schedule_data[slot.day_of_week][slot.time_slot.id].append(slot)

        assigned_subjects = Subject.objects.filter(groups=selected_group).select_related('teacher__user').distinct()

        for subject in assigned_subjects:
            needed = subject.get_weekly_slots_needed()
            
            for l_type, label, color in [('LECTURE', _('Лекция'), 'primary'), ('PRACTICE', _('Практика'), 'success'), ('LAB', _('Лабораторная'), 'info'), ('SRSP', _('СРСП'), 'warning')]:
                needed_weekly = needed.get(l_type, 0)
                
                if needed_weekly > 0:
                    type_slots = [s for s in slots if s.subject_id == subject.id and s.lesson_type == l_type]
                    scheduled_weekly = 0
                    
                    for s in type_slots:
                        scheduled_weekly += 1 if s.week_type == 'EVERY' else 0.5

                    remaining_pairs = max(0, needed_weekly - scheduled_weekly)
                    if remaining_pairs > 0:
                        subjects_to_schedule.append({
                            'obj': subject,
                            'type': l_type,
                            'label': label,
                            'remaining': remaining_pairs,
                            'color': color
                        })

    time_slots = []
    if selected_group:
        try:
            if selected_group.specialty:
                institute = selected_group.specialty.department.faculty.institute
            elif hasattr(request.user, 'dean_profile'):
                institute = request.user.dean_profile.faculty.institute
            else:
                institute = Institute.objects.first()

            shift = active_semester.shift
            time_slots = get_time_slots_for_shift(shift, institute)
        except AttributeError:
            time_slots = TimeSlot.objects.none()
    else:
        time_slots = TimeSlot.objects.none()

    days = [(0, _('Понедельник')), (1, _('Вторник')), (2, _('Среда')), (3, _('Четверг')), (4, _('Пятница')), (5, _('Суббота'))]
    all_semesters = Semester.objects.all().order_by('-is_active', '-start_date')
    if hasattr(request.user, 'dean_profile') or hasattr(request.user, 'vicedean_profile'):
        profile = getattr(request.user, 'dean_profile', None) or getattr(request.user, 'vicedean_profile', None)
        all_semesters = all_semesters.filter(Q(faculty=profile.faculty) | Q(faculty__isnull=True))
    elif hasattr(request.user, 'director_profile') or hasattr(request.user, 'prorector_profile'):
        profile = getattr(request.user, 'director_profile', None) or getattr(request.user, 'prorector_profile', None)
        if profile.institute:
            all_semesters = all_semesters.filter(Q(faculty__institute=profile.institute) | Q(faculty__isnull=True))

    classrooms = Classroom.objects.filter(is_active=True).select_related('building').order_by('building__name', 'number')

    occupied_rooms = {}
    if active_semester and time_slots:
        all_active_slots = ScheduleSlot.objects.filter(
            semester__is_active=True, is_active=True
        ).exclude(classroom__isnull=True).values(
            'day_of_week', 'start_time', 'end_time', 'classroom_id', 'stream_id', 'id', 'week_type'
        )

        for s in all_active_slots:
            d = s['day_of_week']
            rm = s['classroom_id']
            st_id = str(s['stream_id']) if s['stream_id'] else 'None'
            slot_id = str(s['id'])
            wt = s['week_type']

            for ts in time_slots:
                if s['start_time'] < ts.end_time and s['end_time'] > ts.start_time:
                    if d not in occupied_rooms:
                        occupied_rooms[d] = {}
                    if ts.id not in occupied_rooms[d]:
                        occupied_rooms[d][ts.id] = {}
                    if rm not in occupied_rooms[d][ts.id]:
                        occupied_rooms[d][ts.id][rm] = []

                    occupied_rooms[d][ts.id][rm].append({
                        'slot_id': slot_id,
                        'stream_id': st_id,
                        'week_type': wt
                    })

    military_days = set()
    if selected_group and active_semester:
        mil_day_qs = ScheduleSlot.objects.filter(
            group=selected_group,
            semester=active_semester,
            is_military=True,
            is_active=True,
        ).values_list('day_of_week', flat=True).distinct()
        military_days = set(mil_day_qs)
 
    return render(request, 'schedule/constructor_with_limits.html', {
        'groups': groups,
        'group': selected_group,
        'semester': active_semester,
        'all_semesters': all_semesters,
        'time_slots': time_slots,
        'days': days,
        'schedule_data': schedule_data,
        'subjects_to_schedule': subjects_to_schedule,
        'classrooms': classrooms,
        'occupied_rooms': occupied_rooms,
        'active_plan': active_plan,
        'military_days': military_days,  
    })


@login_required
@user_passes_test(is_dean_or_admin)
@require_POST
def create_schedule_slot(request):
    logger.info(f"--- СОЗДАНИЕ СЛОТА РАСПИСАНИЯ: Пользователь {request.user.username} ---")
    try:
        data = json.loads(request.body)
        logger.info(f"Входящие данные: {data}")

        if data.get('is_military_day'):
            try:
                group_id = data.get('group')
                day_of_week = int(data.get('day_of_week'))
                semester_id = data.get('semester_id')

                group = get_object_or_404(Group, id=group_id)
                semester = get_object_or_404(Semester, id=semester_id)

                institute = None
                try:
                    if group.specialty and group.specialty.department.faculty:
                        institute = group.specialty.department.faculty.institute
                except AttributeError:
                    pass

                time_slots = get_time_slots_for_shift(semester.shift, institute)
                if not time_slots.exists():
                    time_slots = get_time_slots_for_shift(semester.shift)
                if not time_slots.exists():
                    return JsonResponse({'success': False, 'error': 'Сначала создайте сетку звонков (Временные слоты) для этой смены в настройках.'}, status=400)

                military_dept = None
                if group.specialty and group.specialty.department:
                    military_dept = group.specialty.department
                else:
                    military_dept = Department.objects.first()

                if not military_dept:
                    return JsonResponse({'success': False, 'error': 'В системе нет ни одной кафедры. Создайте кафедру.'}, status=400)

                military_subject, created_mil = Subject.objects.get_or_create(
                    code="MILITARY",
                    defaults={
                        'name': "Кафедраи ҳарбӣ (Военная кафедра)",
                        'department': military_dept,
                        'type': 'PRACTICE'
                    }
                )

                ScheduleSlot.objects.filter(
                    group=group, semester=semester, day_of_week=day_of_week
                ).delete()

                created_count = 0
                with transaction.atomic():
                    for ts in time_slots:
                        ScheduleSlot.objects.create(
                            group=group,
                            subject=military_subject,
                            semester=semester,
                            day_of_week=day_of_week,
                            time_slot=ts,
                            start_time=ts.start_time,
                            end_time=ts.end_time,
                            is_military=True,
                            lesson_type='PRACTICE',
                            classroom=None,
                            room="Воен. каф",
                            is_active=True
                        )
                        created_count += 1

                return JsonResponse({'success': True, 'message': f'Назначено {created_count} часов военной кафедры'})
            except Exception as e:
                logger.exception(f"create_schedule_slot military: {str(e)}; data={data}")
                return JsonResponse({'success': False, 'error': _('Внутренняя ошибка сервера')}, status=500)

        force = data.get('force', False)
        group_id = data.get('group')
        subject_id = data.get('subject')
        day_of_week = data.get('day_of_week')
        time_slot_id = data.get('time_slot')
        lesson_type = data.get('lesson_type', 'LECTURE')
        semester_id = data.get('semester_id')

        main_group = get_object_or_404(Group, id=group_id)
        subject = get_object_or_404(Subject, id=subject_id)
        time_slot = get_object_or_404(TimeSlot, id=time_slot_id)

        if semester_id:
            active_semester = get_object_or_404(Semester, id=semester_id)
        else:
            active_semester = get_active_semester_for_group(main_group)

        if not active_semester:
            return JsonResponse({'success': False, 'error': _('Нет активного семестра')}, status=400)

        start_h = time_slot.start_time.hour
        if active_semester.shift == 'MORNING' and start_h >= 13:
            return JsonResponse({'success': False, 'error': _('Ошибка смены: Слот ДНЕВНОЙ, а семестр УТРЕННИЙ.')}, status=400)
        if active_semester.shift == 'DAY' and start_h < 13:
            return JsonResponse({'success': False, 'error': _('Ошибка смены: Слот УТРЕННИЙ, а семестр ДНЕВНОЙ.')}, status=400)

        groups_to_schedule = [main_group]
        is_stream = False
        stream_id = None

        if subject.is_stream_subject and subject.groups.count() > 1:
            all_subject_groups = list(subject.groups.all())
            if len(all_subject_groups) > 5:
                return JsonResponse({'success': False, 'error': _('Слишком большой поток! Максимум 5 групп.')}, status=400)
            groups_to_schedule = all_subject_groups
            is_stream = True
            stream_id = uuid.uuid4()

        week_type = data.get('week_type', 'EVERY')

        overlapping_slots = ScheduleSlot.objects.filter(
            semester=active_semester,
            day_of_week=day_of_week,
            is_active=True,
            start_time__lt=time_slot.end_time,
            end_time__gt=time_slot.start_time
        )

        if week_type != 'EVERY':
            overlapping_slots = overlapping_slots.filter(Q(week_type='EVERY') | Q(week_type=week_type))

        conflicts = []
        if not force:
            for target_group in groups_to_schedule:
                busy_group = overlapping_slots.filter(group=target_group).first()
                if busy_group:
                    conflicts.append(f"❌ Группа <b>{busy_group.group.name}</b> уже занята: {busy_group.subject.name} ({busy_group.start_time.strftime('%H:%M')}-{busy_group.end_time.strftime('%H:%M')})")

            if subject.teacher:
                busy_teacher = overlapping_slots.filter(teacher=subject.teacher).first()
                if busy_teacher:
                    conflicts.append(f"❌ Преподаватель <b>{subject.teacher.user.get_full_name()}</b> занят ({busy_teacher.start_time.strftime('%H:%M')}-{busy_teacher.end_time.strftime('%H:%M')})")

        if conflicts:
            return JsonResponse({
                'success': False,
                'is_conflict': True,
                'error': "<br>".join(conflicts)
            }, status=400)

        with transaction.atomic():
            created_slots = []
            for target_group in groups_to_schedule:
                if force:
                    qs_del = ScheduleSlot.objects.filter(
                        group=target_group, day_of_week=day_of_week, time_slot=time_slot, semester=active_semester
                    )
                    if week_type != 'EVERY':
                        qs_del = qs_del.filter(Q(week_type='EVERY') | Q(week_type=week_type))
                    qs_del.delete()

                new_slot = ScheduleSlot.objects.create(
                    group=target_group,
                    subject=subject,
                    teacher=subject.teacher,
                    semester=active_semester,
                    day_of_week=day_of_week,
                    time_slot=time_slot,
                    lesson_type=lesson_type,
                    week_type=week_type,
                    start_time=time_slot.start_time,
                    end_time=time_slot.end_time,
                    stream_id=stream_id if is_stream else None
                )
                created_slots.append(new_slot)

        return JsonResponse({'success': True, 'count': len(created_slots), 'is_stream': is_stream})

    except Exception as e:
        logger.exception(f"ОШИБКА в create_schedule_slot: {str(e)}")
        return JsonResponse({'success': False, 'error': _('Внутренняя ошибка сервера')}, status=500)


@login_required
@require_POST
def update_schedule_room(request, slot_id):
    try:
        data = json.loads(request.body)
        room_id = data.get('room_id')
        room_text = data.get('room', '').strip()
        force = data.get('force', False)
        slot = get_object_or_404(ScheduleSlot, id=slot_id)

        if not room_text and not room_id:
            with transaction.atomic():
                if slot.stream_id:
                    ScheduleSlot.objects.filter(stream_id=slot.stream_id).update(
                        room="",
                        classroom=None
                    )
                else:
                    slot.room = ""
                    slot.classroom = None
                    slot.save()
            return JsonResponse({'success': True, 'room': ''})

        classroom = None
        if room_id and str(room_id).isdigit():
            classroom = Classroom.objects.filter(id=room_id, is_active=True).first()
        elif room_text:
            candidates = Classroom.objects.filter(number__iexact=room_text, is_active=True)
            if candidates.count() == 1:
                classroom = candidates.first()
            elif candidates.count() > 1:
                try:
                    group_institute = slot.group.specialty.department.faculty.institute
                    classroom = candidates.filter(building__institute=group_institute).first()
                except Exception:
                    pass
                if not classroom:
                    classroom = candidates.first()

        if room_text and not classroom:
            return JsonResponse(
                {
                    'success': False,
                    'error': _('Аудитория "%(room)s" не найдена в базе данных. Выберите из списка.') % {'room': room_text}
                },
                status=400
            )

        room_number = classroom.number if classroom else room_text

        occupants = ScheduleSlot.objects.filter(
            semester=slot.semester,
            day_of_week=slot.day_of_week,
            is_active=True,
            start_time__lt=slot.end_time,
            end_time__gt=slot.start_time
        )

        if classroom:
            occupants = occupants.filter(classroom=classroom)
        elif room_number:
            occupants = occupants.filter(room=room_number).exclude(room="")
        else:
            occupants = occupants.none()

        if slot.stream_id:
            occupants = occupants.exclude(stream_id=slot.stream_id)
        else:
            occupants = occupants.exclude(id=slot.id)

        if slot.week_type != 'EVERY':
            occupants = occupants.filter(Q(week_type='EVERY') | Q(week_type=slot.week_type))

        if occupants.exists() and not force:
            other = occupants.first()
            b_name = classroom.building.name if classroom and classroom.building else "Корпус"
            r_num = classroom.number if classroom else room_number
            return JsonResponse({
                'success': False,
                'is_conflict': True,
                'error': _('Конфликт в %(building)s, каб. %(num)s! Там уже сидит: %(group)s (%(subject)s)') % {
                    'building': b_name,
                    'num': r_num,
                    'group': other.group.name,
                    'subject': other.subject.name
                }
            }, status=400)

        if not force and classroom:
            
            room_types_dict = dict(ROOM_TYPES)

            if slot.subject.preferred_room_type and classroom.room_type != slot.subject.preferred_room_type:
                if not (slot.lesson_type == 'LECTURE' and classroom.room_type == 'LECTURE'):
                    return JsonResponse({
                        'success': False,
                        'is_type_warning': True,
                        'error': _('⚠️ ВНИМАНИЕ: Для предмета "%(subject)s" рекомендован тип кабинета "%(pref)s", а вы выбрали "%(curr)s".') % {
                            'subject': slot.subject.name,
                            'pref': room_types_dict.get(slot.subject.preferred_room_type, ''),
                            'curr': classroom.get_room_type_display()
                        }
                    }, status=400)

            elif slot.lesson_type == 'LECTURE' and classroom.room_type not in ['LECTURE', 'SPORT']:
                return JsonResponse({
                    'success': False,
                    'is_type_warning': True,
                    'error': _('⚠️ ВНИМАНИЕ: Вы назначаете ЛЕКЦИЮ в кабинет типа "%(type)s". Обычно лекции проводятся в больших Лекционных аудиториях.') % {'type': classroom.get_room_type_display()}
                }, status=400)
            elif slot.lesson_type in ['PRACTICE', 'SRSP'] and classroom.room_type == 'LECTURE':
                return JsonResponse({
                    'success': False,
                    'is_type_warning': True,
                    'error': _('⚠️ ВНИМАНИЕ: Вы назначаете ПРАКТИКУ в большую ЛЕКЦИОННУЮ аудиторию. Это неэффективное использование пространства.')
                }, status=400)

        with transaction.atomic():
            if slot.stream_id:
                ScheduleSlot.objects.filter(stream_id=slot.stream_id).update(
                    room=room_number,
                    classroom=classroom
                )
            else:
                slot.room = room_number
                slot.classroom = classroom
                slot.save()

        display_name = room_number
        return JsonResponse({'success': True, 'room': display_name})

    except Exception as e:
        logger.exception("update_schedule_room")
        return JsonResponse({'success': False, 'error': _('Внутренняя ошибка сервера')}, status=500)

@login_required
@require_POST
def clear_schedule(request):
    if not is_dean_or_admin(request.user):
        return JsonResponse({'success': False, 'error': _('Нет прав')}, status=403)
        
    try:
        data = json.loads(request.body)
        group_id = data.get('group_id')
        semester_id = data.get('semester_id')
        
        group = get_object_or_404(Group, id=group_id)
        semester = get_object_or_404(Semester, id=semester_id)
        
        if hasattr(request.user, 'dean_profile'):
            faculty = request.user.dean_profile.faculty
            if group.specialty and group.specialty.department.faculty != faculty:
                return JsonResponse({'success': False, 'error': _('Нет доступа к этой группе')}, status=403)
                
        slots = ScheduleSlot.objects.filter(group=group, semester=semester)
        slots.update(is_active=False)
        slots.delete()
        return JsonResponse({'success': True})
    except Exception as e:
        logger.exception("clear_schedule")
        return JsonResponse({'success': False, 'error': _('Внутренняя ошибка сервера')}, status=500)

@login_required
@require_POST
def delete_schedule_slot(request, slot_id):
    try:
        schedule_slot = get_object_or_404(ScheduleSlot, id=slot_id)

        if not (request.user.is_staff or hasattr(request.user, 'dean_profile')):
            return JsonResponse({'success': False, 'error': _('Нет прав')}, status=403)

        if hasattr(request.user, 'dean_profile'):
            faculty = request.user.dean_profile.faculty
            if schedule_slot.group.specialty and schedule_slot.group.specialty.department.faculty != faculty:
                return JsonResponse({'success': False, 'error': _('Нет доступа к этому занятию')}, status=403)

        with transaction.atomic():
            if schedule_slot.stream_id:
                count = ScheduleSlot.objects.filter(stream_id=schedule_slot.stream_id).delete()[0]
                msg = _('Удален поток (%(count)s групп)') % {'count': count}
            else:
                schedule_slot.delete()
                msg = _('Занятие удалено')

        return JsonResponse({'success': True, 'message': msg})

    except Exception as e:
        logger.exception("delete_schedule_slot")
        return JsonResponse({'success': False, 'error': _('Внутренняя ошибка сервера')}, status=500)

@login_required
def schedule_view(request):
    user = request.user
    group = None
    teacher = None
    active_semester = None
    
    if hasattr(user, 'student_profile'):
        try:
            student = user.student_profile
            group = student.group
            if group:
                active_semester = get_active_semester_for_group(group)
        except Student.DoesNotExist:
            pass

    if not active_semester:
        active_semester = Semester.objects.filter(is_active=True).first()

    if not active_semester:
        active_semester = Semester.objects.order_by('-start_date').first()
        if not active_semester:
             messages.warning(request, _('В системе не создано ни одного семестра.'))
             return render(request, 'schedule/no_semester.html')

    groups = Group.objects.none()
    is_management = hasattr(user, 'dean_profile') or user.is_superuser or hasattr(user, 'director_profile') or hasattr(user, 'prorector_profile')
    is_teacher = hasattr(user, 'teacher_profile')
    
    if is_management:
        if 'group' in request.GET:
            group_id = request.GET.get('group')
            if group_id:
                request.session['last_schedule_group'] = group_id
            else:
                request.session.pop('last_schedule_group', None)
        else:
            group_id = request.session.get('last_schedule_group')
        
        if user.is_superuser or hasattr(user, 'director_profile') or hasattr(user, 'prorector_profile'):
            institutes = Institute.objects.all()
            
            if 'institute' in request.GET:
                institute_id = request.GET.get('institute')
                if institute_id:
                    request.session['last_schedule_institute'] = institute_id
                else:
                    request.session.pop('last_schedule_institute', None)
            else:
                institute_id = request.session.get('last_schedule_institute')
                
            groups = Group.objects.all().select_related('specialty__department__faculty__institute')
            if institute_id:
                groups = groups.filter(specialty__department__faculty__institute_id=institute_id)
        else:
            groups = Group.objects.filter(specialty__department__faculty=user.dean_profile.faculty)

        if group_id:
            group = Group.objects.filter(id=group_id).first()
            if not group:
                request.session.pop('last_schedule_group', None)
            else:
                active_semester = get_active_semester_for_group(group)
                if not active_semester:
                    messages.warning(request, _('Нет активного семестра для %(course)s курса') % {'course': group.course})
        elif is_teacher and request.GET.get('view') != 'groups':
            teacher = user.teacher_profile
            
    elif is_teacher:
        teacher = user.teacher_profile
        group_ids = ScheduleSlot.objects.filter(
            teacher=teacher, semester=active_semester, is_active=True
        ).values_list('group_id', flat=True).distinct()

        groups = Group.objects.filter(id__in=group_ids)
        group_id = request.GET.get('group')

        if group_id:
            group = get_object_or_404(Group, id=group_id, id__in=group_ids)

    slots = None
    target_key = None
    time_slots = TimeSlot.objects.none()
    
    if group and active_semester:
        institute = None
        if group.specialty and group.specialty.department.faculty:
            institute = group.specialty.department.faculty.institute
        time_slots = get_time_slots_for_shift(active_semester.shift, institute)
        valid_slot_ids = list(time_slots.values_list('id', flat=True))

        slots = ScheduleSlot.objects.filter(
            group=group,
            semester=active_semester,
            time_slot_id__in=valid_slot_ids,
            is_active=True
        ).select_related('subject', 'teacher__user', 'time_slot')
        target_key = group.id

    elif teacher and active_semester and not group:
        used_slot_ids = ScheduleSlot.objects.filter(
            teacher=teacher, semester=active_semester, is_active=True
        ).values_list('time_slot_id', flat=True).distinct()
        
        if used_slot_ids:
            time_slots = TimeSlot.objects.filter(id__in=used_slot_ids).order_by('start_time')
        else:
            institute = None
            if teacher.department and teacher.department.faculty:
                institute = teacher.department.faculty.institute
            time_slots = get_time_slots_for_shift(active_semester.shift, institute)

        slots = ScheduleSlot.objects.filter(
            teacher=teacher,
            semester=active_semester,
            is_active=True
        ).select_related('subject', 'group', 'time_slot')
        target_key = 'teacher'

    if slots is not None:
        days =[(0, _('ДУШАНБЕ')), (1, _('СЕШАНБЕ')), (2, _('ЧОРШАНБЕ')), (3, _('ПАНҶШАНБЕ')), (4, _('ҶУМЪА')), (5, _('ШАНБЕ'))]
        
        schedule_data = {target_key: {}}
        for slot in slots:
            if slot.day_of_week not in schedule_data[target_key]:
                schedule_data[target_key][slot.day_of_week] = {}
            if slot.time_slot.id not in schedule_data[target_key][slot.day_of_week]:
                schedule_data[target_key][slot.day_of_week][slot.time_slot.id] =[]
            schedule_data[target_key][slot.day_of_week][slot.time_slot.id].append(slot)

        context = {
            'group': group,
            'is_teacher_view': (target_key == 'teacher'),
            'target_key': target_key,
            'groups': groups if (hasattr(user, 'dean_profile') or user.is_superuser or hasattr(user, 'director_profile') or hasattr(user, 'prorector_profile') or hasattr(user, 'teacher_profile')) else None,
            'institutes': locals().get('institutes', None),
            'days': days,
            'time_slots': time_slots,
            'schedule_data': schedule_data,
            'active_semester': active_semester,
            'is_view_mode': True,
        }
        return render(request, 'schedule/schedule_view_unified.html', context)

    context = {
        'group': group,
        'groups': groups if (hasattr(user, 'dean_profile') or user.is_superuser or hasattr(user, 'director_profile') or hasattr(user, 'prorector_profile') or hasattr(user, 'teacher_profile')) else None,
        'institutes': locals().get('institutes', None),
        'active_semester': active_semester,
    }
    return render(request, 'schedule/schedule_view_unified.html', context)

@login_required
def today_classes(request):
    user = request.user
    today = datetime.now()
    day_of_week = today.weekday()
    current_time = today.time()

    if hasattr(user, 'student_profile'):
        try:
            student = user.student_profile
            if student.group:
                active_semester = get_active_semester_for_group(student.group)
            else:
                active_semester = None
        except Exception:
            logger.exception("today_classes active_semester")
            active_semester = None
    else:
        active_semester = Semester.objects.filter(is_active=True).first()

    classes = []
    if not active_semester:
        return render(request, 'schedule/today_widget.html', {'classes': classes, 'current_time': current_time, 'today': today})

    if hasattr(user, 'student_profile'):
        try:
            student = user.student_profile
            if student.group:
                classes = ScheduleSlot.objects.filter(
                    group=student.group, semester=active_semester,
                    day_of_week=day_of_week, is_active=True
                ).select_related('subject', 'teacher__user').order_by('start_time')
        except Student.DoesNotExist:
            pass

    elif hasattr(user, 'teacher_profile'):
        try:
            teacher = user.teacher_profile
            classes = ScheduleSlot.objects.filter(
                teacher=teacher, semester=active_semester,
                day_of_week=day_of_week, is_active=True
            ).select_related('subject', 'group').order_by('start_time')
        except Teacher.DoesNotExist:
            pass

    return render(request, 'schedule/today_widget.html', {
        'classes': classes, 'current_time': current_time, 'today': today
    })


@login_required
@user_passes_test(is_dean_or_admin)
def manage_subjects(request):
    subjects = Subject.objects.all().select_related('teacher__user', 'department')
    
    if hasattr(request.user, 'dean_profile'):
        subjects = subjects.filter(department__faculty=request.user.dean_profile.faculty)
        
    search = request.GET.get('search', '')
    if search:
        subjects = subjects.filter(Q(name__icontains=search) | Q(code__icontains=search))
        
    return render(request, 'schedule/manage_subjects.html', {'subjects': subjects, 'search': search})



@login_required
@user_passes_test(is_dean_or_admin)
@login_required
@user_passes_test(is_dean_or_admin)
def add_subject(request):
    initial_data = {}
    dept_id = request.GET.get('department')

    target_department = None
    if dept_id:
        target_department = get_object_or_404(Department, id=dept_id)
        if hasattr(request.user, 'dean_profile'):
            if target_department.faculty != request.user.dean_profile.faculty:
                messages.error(request, _("Нет доступа к этой кафедре"))
                return redirect('schedule:manage_subjects')

        initial_data['department'] = target_department
        initial_data['assign_to_all_groups'] = True

    if request.method == 'POST':
        form = SubjectForm(request.POST)
        if form.is_valid():
            subject = form.save(commit=False)

            if target_department and not subject.department_id:
                subject.department = target_department

            assign_to_all = form.cleaned_data.get('assign_to_all_groups')
            groups = form.cleaned_data.get('groups')
            
            if assign_to_all and subject.department:
                groups = Group.objects.filter(specialty__department=subject.department)
                
            if not groups:
                subject.save()
                form.save_m2m()
                messages.success(request, _('Предмет "%(subject_name)s" создан') % {'subject_name': subject.name})
            else:
                if subject.is_stream_subject:
                    subject.save()
                    subject.groups.set(groups)
                    subject.required_competencies.set(form.cleaned_data.get('required_competencies',[]))
                    messages.success(request, _('Потоковый предмет "%(subject_name)s" создан') % {'subject_name': subject.name})
                else:
                    created_count = 0
                    first = True
                    for grp in groups:
                        if first:
                            subject.save()
                            subject.groups.add(grp)
                            subject.required_competencies.set(form.cleaned_data.get('required_competencies',[]))
                            first = False
                            created_count += 1
                        else:
                            short_uuid = uuid.uuid4().hex[:6]
                            new_subject = Subject.objects.create(
                                name=subject.name,
                                code=f"{subject.code}_{short_uuid}",
                                department=subject.department,
                                type=subject.type,
                                lecture_hours=subject.lecture_hours,
                                practice_hours=subject.practice_hours,
                                lab_hours=subject.lab_hours,
                                control_hours=subject.control_hours,
                                independent_work_hours=subject.independent_work_hours,
                                semester_weeks=subject.semester_weeks,
                                is_stream_subject=False,
                                preferred_room_type=subject.preferred_room_type,
                                teacher=subject.teacher,
                                description=subject.description,
                                credits=subject.credits,
                                credit_type=subject.credit_type,
                                plan_discipline=subject.plan_discipline
                            )
                            new_subject.groups.add(grp)
                            new_subject.required_competencies.set(form.cleaned_data.get('required_competencies',[]))
                            created_count += 1
                    messages.success(request, _('Создано %(count)s отдельных предметов "%(subject_name)s"') % {'count': created_count, 'subject_name': subject.name})

            if dept_id:
                return redirect('accounts:manage_structure')
            return redirect('schedule:manage_subjects')
        else:
            print(_("Ошибки формы предмета:"), form.errors)
    else:
        form = SubjectForm(initial=initial_data)

        if target_department:
            form.fields['teacher'].queryset = Teacher.objects.filter(
                Q(department=target_department) | Q(additional_departments=target_department)
            ).distinct()
            form.fields['department'].queryset = Department.objects.filter(id=target_department.id)
        elif hasattr(request.user, 'dean_profile'):
            faculty = request.user.dean_profile.faculty
            form.fields['department'].queryset = Department.objects.filter(faculty=faculty)
            form.fields['teacher'].queryset = Teacher.objects.filter(
                Q(department__faculty=faculty) | 
                Q(additional_departments__faculty=faculty) |
                Q(subject__department__faculty=faculty) |
                Q(subject__groups__specialty__department__faculty=faculty) |
                Q(scheduleslot__group__specialty__department__faculty=faculty)
            ).distinct()

    faculty = request.user.dean_profile.faculty if hasattr(request.user, 'dean_profile') else None
    credit_templates = CreditTemplate.objects.filter(Q(faculty=faculty) | Q(faculty__isnull=True)).order_by('credits')

    inst = None
    if target_department and getattr(target_department, 'faculty', None):
        inst = target_department.faculty.institute
    elif faculty:
        inst = faculty.institute

    return render(request, 'schedule/subject_form.html', {
        'form': form,
        'title': _('Добавить предмет'),
        'credit_templates': credit_templates,
        'institute_pair_ratio': institute_pair_ratio(inst),
    })


@login_required
@user_passes_test(is_dean_or_admin)
@require_POST
def split_subject(request, subject_id):
    subject = get_object_or_404(Subject, id=subject_id)
    
    if subject.is_stream_subject or subject.groups.count() <= 1:
        messages.error(request, "Этот предмет не нуждается в разделении.")
        return redirect('schedule:manage_subjects')
        
    groups = list(subject.groups.all())
    
    with transaction.atomic():
        first_group = groups[0]
        subject.groups.clear()
        subject.groups.add(first_group)
        
        for grp in groups[1:]:
            short_uuid = uuid.uuid4().hex[:6]
            new_subject = Subject.objects.create(
                name=subject.name,
                code=f"{subject.code}_{short_uuid}",
                department=subject.department,
                type=subject.type,
                lecture_hours=subject.lecture_hours,
                practice_hours=subject.practice_hours,
                lab_hours=subject.lab_hours,
                control_hours=subject.control_hours,
                independent_work_hours=subject.independent_work_hours,
                semester_weeks=subject.semester_weeks,
                is_stream_subject=False,
                preferred_room_type=subject.preferred_room_type,
                teacher=subject.teacher,
                description=subject.description,
                credits=subject.credits,
                credit_type=subject.credit_type,
                plan_discipline=subject.plan_discipline
            )
            new_subject.groups.add(grp)
            new_subject.required_competencies.set(subject.required_competencies.all())
            
            ScheduleSlot.objects.filter(subject=subject, group=grp).update(subject=new_subject)
            
            from journal.models import JournalEntry, StudentMatrixScore
            JournalEntry.objects.filter(subject=subject, student__group=grp).update(subject=new_subject)
            StudentMatrixScore.objects.filter(subject=subject, student__group=grp).update(subject=new_subject)
            
    messages.success(request, f"Предмет успешно разделен на {len(groups)} отдельных предметов!")
    return redirect('schedule:manage_subjects')


@login_required
@user_passes_test(is_dean_or_admin)
def edit_subject(request, subject_id):
    subject = get_object_or_404(Subject, id=subject_id)
    if hasattr(request.user, 'dean_profile'):
        if subject.department.faculty != request.user.dean_profile.faculty:
            messages.error(request, _("Нет доступа к этому предмету"))
            return redirect('schedule:manage_subjects')
            
    if request.method == 'POST':
        old_groups = list(subject.groups.all()) 
        form = SubjectForm(request.POST, instance=subject)
        
        if hasattr(request.user, 'dean_profile'):
            faculty = request.user.dean_profile.faculty
            form.fields['department'].queryset = Department.objects.filter(faculty=faculty)
            form.fields['teacher'].queryset = Teacher.objects.filter(
                Q(department__faculty=faculty) | 
                Q(additional_departments__faculty=faculty) |
                Q(subject__department__faculty=faculty) |
                Q(subject__groups__specialty__department__faculty=faculty) |
                Q(scheduleslot__group__specialty__department__faculty=faculty)
            ).distinct()
            
        if form.is_valid():
            saved_subject = form.save(commit=False)
            saved_subject.save()
            
            submitted_groups = form.cleaned_data.get('groups')
            if submitted_groups and submitted_groups.count() > 0:
                form.save_m2m() 
            else:
                saved_subject.groups.set(old_groups) 
                
            messages.success(request, _('Предмет обновлен'))
            return redirect('schedule:manage_subjects')
    else:
        form = SubjectForm(instance=subject)
        if hasattr(request.user, 'dean_profile'):
            faculty = request.user.dean_profile.faculty
            form.fields['department'].queryset = Department.objects.filter(faculty=faculty)
            form.fields['teacher'].queryset = Teacher.objects.filter(
                Q(department__faculty=faculty) | 
                Q(additional_departments__faculty=faculty) |
                Q(subject__department__faculty=faculty) |
                Q(subject__groups__specialty__department__faculty=faculty) |
                Q(scheduleslot__group__specialty__department__faculty=faculty)
            ).distinct()
            
    faculty = request.user.dean_profile.faculty if hasattr(request.user, 'dean_profile') else None
    credit_templates = CreditTemplate.objects.filter(Q(faculty=faculty) | Q(faculty__isnull=True)).order_by('credits')

    inst = subject.department.faculty.institute if subject.department and subject.department.faculty else None

    return render(request, 'schedule/subject_form.html', {
        'form': form,
        'subject': subject,
        'title': _('Редактировать предмет'),
        'credit_templates': credit_templates,
        'institute_pair_ratio': institute_pair_ratio(inst),
    })

@login_required
@user_passes_test(is_dean_or_admin)
def delete_subject(request, subject_id):
    subject = get_object_or_404(Subject, id=subject_id)
    if hasattr(request.user, 'dean_profile'):
        if subject.department.faculty != request.user.dean_profile.faculty:
            messages.error(request, _("Нет доступа к этому предмету"))
            return redirect('schedule:manage_subjects')

    subject.delete()
    messages.success(request, _('Предмет удален'))
    return redirect('schedule:manage_subjects')
@login_required 
@user_passes_test(is_dean_or_admin)
def manage_semesters(request):
    semesters = Semester.objects.all().order_by('-academic_year', 'course')
    
    if hasattr(request.user, 'dean_profile'):
        faculty = request.user.dean_profile.faculty
        semesters = semesters.filter(faculty=faculty)

    return render(request, 'schedule/manage_semesters.html', {
        'semesters': semesters
    })

@user_passes_test(is_dean_or_admin)
def add_semester(request):
    if request.method == 'POST':
        form = SemesterForm(request.POST, user=request.user)
        if form.is_valid():
            semester = form.save()
            faculty_name = semester.faculty.name if semester.faculty else _("Всего университета")
            messages.success(request, _('Семестр "%(semester_name)s" создан (%(faculty_name)s)') % {
                'semester_name': semester.name, 
                'faculty_name': faculty_name
            })
            return redirect('schedule:manage_semesters')
    else:
        form = SemesterForm(user=request.user)
    return render(request, 'schedule/semester_form.html', {'form': form, 'title': _('Добавить семестр')})

@user_passes_test(is_dean_or_admin)
def edit_semester(request, semester_id):
    semester = get_object_or_404(Semester, id=semester_id)
    if request.method == 'POST':
        form = SemesterForm(request.POST, instance=semester, user=request.user)
        if form.is_valid():
            semester = form.save(commit=False)
            semester.save()
            
            dept = form.cleaned_data.get('department_filter')
            course = form.cleaned_data.get('course')
            
            if dept and course:
                dept_groups = Group.objects.filter(specialty__department=dept, course=course)
                semester.groups.add(*dept_groups)
                
            form.save_m2m()
            messages.success(request, _('Семестр обновлен'))
            return redirect('schedule:manage_semesters')
    else:
        form = SemesterForm(instance=semester, user=request.user)
        
    return render(request, 'schedule/semester_form.html', {'form': form, 'semester': semester, 'title': _('Редактировать семестр')})

@user_passes_test(is_dean_or_admin)
def toggle_semester_active(request, semester_id):
    semester = get_object_or_404(Semester, id=semester_id)
    
    with transaction.atomic():
        if not semester.is_active:
            Semester.objects.filter(
                faculty=semester.faculty,
                course=semester.course, 
                is_active=True
            ).update(is_active=False)
            
            semester.is_active = True
            semester.save()
            messages.success(request, _('Семестр "%(semester_name)s" (%(course)s курс) теперь АКТИВЕН!') % {'semester_name': semester.name, 'course': semester.course})
        else:
            semester.is_active = False
            semester.save()
            messages.warning(request, _('Семестр "%(semester_name)s" деактивирован.') % {'semester_name': semester.name})

    return redirect('schedule:manage_semesters')

@login_required
def manage_classrooms(request):
    classrooms = Classroom.objects.select_related('building').all().order_by('building', 'floor', 'number')
    
    institutes = Institute.objects.all()
    selected_institute_id = request.GET.get('institute')
    
    if request.user.is_superuser or hasattr(request.user, 'director_profile') or hasattr(request.user, 'prorector_profile'):
        if selected_institute_id:
            classrooms = classrooms.filter(building__institute_id=selected_institute_id)
    
    elif hasattr(request.user, 'dean_profile'):
        faculty = request.user.dean_profile.faculty
        if faculty and faculty.institute:
            classrooms = classrooms.filter(building__institute=faculty.institute)
            institutes = [] 

    return render(request, 'schedule/manage_classrooms.html', {
        'classrooms': classrooms,
        'institutes': institutes,
        'selected_institute_id': int(selected_institute_id) if selected_institute_id else None,
        'is_admin': request.user.is_superuser or hasattr(request.user, 'director_profile') or hasattr(request.user, 'prorector_profile'),
        'is_facility_admin': is_facility_admin(request.user),   
    })

@user_passes_test(is_facility_admin)
def add_classroom(request):
    if request.method == 'POST':
        form = ClassroomForm(request.POST, user=request.user)
        if form.is_valid():
            classroom = form.save()
            messages.success(request, _('Кабинет %(classroom_number)s добавлен') % {'classroom_number': classroom.number})
            return redirect('schedule:manage_classrooms')
    else:
        form = ClassroomForm(user=request.user)
    return render(request, 'schedule/classroom_form.html', {'form': form, 'title': _('Добавить кабинет')})

@user_passes_test(is_facility_admin)
def bulk_add_classrooms(request):
    if request.method == 'POST':
        form = BulkClassroomForm(request.POST, user=request.user)
        if form.is_valid():
            building = form.cleaned_data['building']
            floor = form.cleaned_data['floor']
            start = form.cleaned_data['start_number']
            end = form.cleaned_data['end_number']
            capacity = form.cleaned_data['capacity']
            room_type = form.cleaned_data['room_type']

            created = 0
            for num in range(start, end + 1):
                number = f"{num}"
                if not Classroom.objects.filter(building=building, number=number).exists():
                    Classroom.objects.create(
                        building=building, 
                        number=number, 
                        floor=floor, 
                        capacity=capacity,
                        room_type=room_type 
                    )
                    created += 1

            messages.success(request, _('Создано %(created_count)s кабинетов в корпусе %(building)s') % {
                'created_count': created,
                'building': building.name
            })
            return redirect('schedule:manage_classrooms')
    else:
        form = BulkClassroomForm(user=request.user)
    return render(request, 'schedule/bulk_classroom_form.html', {'form': form})

@user_passes_test(is_facility_admin)
def delete_classroom(request, classroom_id):
    classroom = get_object_or_404(Classroom, id=classroom_id)
    classroom.delete()
    messages.success(request, _('Кабинет %(classroom_number)s удален') % {'classroom_number': classroom.number})
    return redirect('schedule:manage_classrooms')

@login_required
def group_list(request):
    user = request.user

    if hasattr(user, 'teacher_profile'):
        try:
            teacher = user.teacher_profile
            subject_group_ids = Subject.objects.filter(teacher=teacher).values_list('groups__id', flat=True)
            schedule_group_ids = ScheduleSlot.objects.filter(teacher=teacher, is_active=True).values_list('group_id', flat=True)
            
            all_ids = set(filter(None, list(subject_group_ids) + list(schedule_group_ids)))
            groups = Group.objects.filter(id__in=all_ids).distinct().order_by('course', 'name')
        except Teacher.DoesNotExist:
            groups = Group.objects.none()
            
    elif hasattr(user, 'dean_profile'):
        groups = Group.objects.filter(specialty__department__faculty=user.dean_profile.faculty).order_by('course', 'name')
    elif user.is_superuser:
        groups = Group.objects.all().order_by('course', 'name')
    else:
        messages.error(request, _('Доступ запрещен'))
        return redirect('core:dashboard')

    groups_with_students =[]
    for group in groups:
        students = Student.objects.filter(group=group).select_related('user').order_by('user__last_name')
        groups_with_students.append({'group': group, 'students': students})

    return render(request, 'schedule/group_list.html', {'groups_with_students': groups_with_students})

@login_required
def export_schedule(request):
    if not DOCX_AVAILABLE:
        return HttpResponse(_("Library python-docx not installed"), status=500)

    group_id = request.GET.get('group')
    
    current_site_lang = translation.get_language()
    if current_site_lang:
        current_site_lang = current_site_lang[:2]
    else:
        current_site_lang = 'ru'
        
    lang = request.GET.get('lang', current_site_lang)
    
    if lang not in ['ru', 'tg', 'en']:
        lang = 'ru'
        
    translation.activate(lang)
    
    try:
        group = get_object_or_404(Group, id=group_id)
        active_semester = get_active_semester_for_group(group)

        if not active_semester:
            return HttpResponse(_("Нет активного семестра"), status=400)

        try:
            if group.specialty:
                specialty_code = group.specialty.code
                specialty_name = group.specialty.name
                department = group.specialty.department
                faculty = department.faculty
                institute = faculty.institute
            else:
                specialty_code = "—"
                specialty_name = _("Умумитаълимӣ (Общая)")
                faculty = active_semester.faculty
                institute = faculty.institute if faculty else Institute.objects.first()
        except AttributeError:
            return HttpResponse(_("Ошибка структуры: Невозможно определить институт."), status=400)

        director_user = None
        director_obj = Director.objects.filter(institute=institute).first()
        if director_obj:
            director_user = director_obj.user

        vice_user = None
        vice_obj = ProRector.objects.filter(institute=institute, title__icontains='таълим').first()
        if not vice_obj:
            vice_obj = ProRector.objects.filter(institute=institute).first()

        if vice_obj:
            vice_user = vice_obj.user

        director_name = director_user.get_full_name() if director_user else "__________________"
        vice_name = vice_user.get_full_name() if vice_user else "__________________"
        head_edu_obj = ProRector.objects.filter(
            institute=institute,
            title__icontains='раёсат'
        ).exclude(
            title__icontains='директор'
        ).first()

        head_edu_name = head_edu_obj.user.get_full_name() if head_edu_obj else "__________________"

        trans = {
            'ru': {
                'agreed': "Согласовано:",
                'head_edu': "Начальник учебного управления",
                'approved': "Утверждаю:",
                'vice_director': "Заместитель директора\nпо учебной работе",
                'docent': "________ доцент",
                'date_line': "«___» _________ 202__ г.",
                'title': "Расписание уроков",
                'subtitle': "на {sem_text} семестр {year_text} учебного года для студентов {course}-го курса {institute_name}",
                'sem_1': "первый",
                'sem_2': "второй",
                'shift': "(СМЕНА {shift_num})",
                'week': "НЕДЕЛЯ",
                'time': "ЧАС",
                'aud': "АУД",
                'military': "Военная кафедра",
                'director': "Директор",
                'days':['Понедельник', 'Вторник', 'Среда', 'Четверг', 'Пятница', 'Суббота'],
                'students': "чел."
            },
            'tg': {
                'agreed': "Мувофиқа карда шуд:",
                'head_edu': "Сардори раёсати таълим",
                'approved': "Тасдиқ мекунам:",
                'vice_director': "Муовини директор\nоид ба корҳои таълимӣ",
                'docent': "________ дотсент",
                'date_line': "«___» _________ 202__ с.",
                'title': "ҶАДВАЛИ ДАРСӢ",
                'subtitle': "дар нимсолаи {sem_text} соли таҳсили {year_text} барои донишҷӯёни курси {course}-юми {institute_name}",
                'sem_1': "якуми",
                'sem_2': "дуюми",
                'shift': "(БАСТИ {shift_num})",
                'week': "РӮЗҲО",
                'time': "СОАТ",
                'aud': "ҲУҶРА",
                'military': "Кафедраи ҳарбӣ",
                'director': "Директор",
                'days':['Душанбе', 'Сешанбе', 'Чоршанбе', 'Панҷшанбе', 'Ҷумъа', 'Шанбе'],
                'students': "нафар"
            },
            'en': {
                'agreed': "Agreed:",
                'head_edu': "Head of Educational Department",
                'approved': "Approved:",
                'vice_director': "Deputy Director\nfor Academic Affairs",
                'docent': "________ docent",
                'date_line': "«___» _________ 202__",
                'title': "Class Schedule",
                'subtitle': "for the {sem_text} semester of {year_text} academic year for {course} year students of {institute_name}",
                'sem_1': "first",
                'sem_2': "second",
                'shift': "(SHIFT {shift_num})",
                'week': "DAY",
                'time': "TIME",
                'aud': "ROOM",
                'military': "Military Department",
                'director': "Director",
                'days':['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday'],
                'students': "students"
            }
        }
        
        t = trans[lang]

        doc = Document()

        section = doc.sections[0]
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)

        header_table = doc.add_table(rows=1, cols=2)
        header_table.autofit = True
        header_table.width = section.page_width - section.left_margin - section.right_margin

        c1 = header_table.cell(0, 0)
        p1 = c1.paragraphs[0]
        p1.add_run(t['agreed'] + "\n").bold = True
        p1.add_run(t['head_edu'] + "\n")
        p1.add_run(f"{t['docent']} {head_edu_name}\n")
        p1.add_run(t['date_line'])

        c2 = header_table.cell(0, 1)
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p2.add_run(t['approved'] + "\n").bold = True
        p2.add_run(t['vice_director'] + "\n").bold = False
        p2.add_run(f"{t['docent']} {vice_name}\n")
        p2.add_run(t['date_line'])

        doc.add_paragraph()

        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run(t['title'])
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = 'Times New Roman'

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sem_text = t['sem_1'] if active_semester.number == 1 else t['sem_2']
        year_text = active_semester.academic_year

        text = t['subtitle'].format(
            sem_text=sem_text,
            year_text=year_text,
            course=group.course,
            institute_name=institute.name
        )

        run_sub = subtitle.add_run(text)
        run_sub.font.name = 'Times New Roman'
        run_sub.font.size = Pt(12)
        run_sub.bold = True

        shift_num = "1" if active_semester.shift == "MORNING" else "2"
        shift_p = doc.add_paragraph(t['shift'].format(shift_num=shift_num))
        shift_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shift_p.runs[0].bold = True
        shift_p.runs[0].font.size = Pt(14)

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.autofit = False

        def set_col_widths(row):
            row.cells[0].width = Cm(1.5)
            row.cells[1].width = Cm(2.5)
            row.cells[2].width = Cm(11.0)
            row.cells[3].width = Cm(2.0)

        set_col_widths(table.rows[0])

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = t['week']
        hdr_cells[1].text = t['time']

        hdr_cells[2].text = f"Группа {group.name} | {specialty_code} – «{specialty_name}» ({group.students.count()} {t['students']})"
        hdr_cells[3].text = t['aud']

        for cell in hdr_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:val'), 'clear')
            shading_elm.set(qn('w:color'), 'auto')
            shading_elm.set(qn('w:fill'), 'D9D9D9')
            cell._tc.get_or_add_tcPr().append(shading_elm)

            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(10)

        time_slots = get_time_slots_for_shift(active_semester.shift, institute)

        days = list(enumerate(t['days']))

        for day_num, day_name in days:
            is_military_day = ScheduleSlot.objects.filter(
                group=group, semester=active_semester, day_of_week=day_num, is_military=True
            ).exists()

            first_row_idx = len(table.rows)

            for ts in time_slots:
                row = table.add_row()
                set_col_widths(row)
                row.cells[1].text = f'{ts.start_time.strftime("%H:%M")}-{ts.end_time.strftime("%H:%M")}'
                row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                row.cells[1].paragraphs[0].runs[0].font.bold = True

                if not is_military_day:
                    cell_slots = ScheduleSlot.objects.filter(
                        group=group, semester=active_semester, day_of_week=day_num, time_slot=ts, is_active=True
                    )

                    if cell_slots.exists():
                        cell = row.cells[2]
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        rooms =[]
                        for idx, slot in enumerate(cell_slots):
                            lesson_type_display = slot.get_lesson_type_display()
                            week_mark = ""
                            if slot.week_type == 'RED': week_mark = " (Красн.)"
                            elif slot.week_type == 'BLUE': week_mark = " (Син.)"
                            
                            if idx > 0:
                                p.add_run("\n-------------------\n")

                            run = p.add_run(f"{slot.subject.name} ({lesson_type_display}){week_mark}\n")
                            run.bold = True
                            if slot.teacher:
                                p.add_run(f"{slot.teacher.user.get_full_name()}")

                            if slot.room:
                                rooms.append(slot.room)

                        cell_aud = row.cells[3]
                        cell_aud.text = "\n".join(set(rooms)) if rooms else ""
                        cell_aud.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cell_aud.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            day_cell = table.rows[first_row_idx].cells[0]
            day_cell.text = day_name
            day_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            tcPr = day_cell._tc.get_or_add_tcPr()
            textDirection = OxmlElement('w:textDirection')
            textDirection.set(qn('w:val'), 'btLr')
            tcPr.append(textDirection)

            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:val'), 'clear')
            shading_elm.set(qn('w:color'), 'auto')
            shading_elm.set(qn('w:fill'), 'D9D9D9')
            tcPr.append(shading_elm)

            last_row_idx = len(table.rows) - 1
            if last_row_idx > first_row_idx:
                day_cell.merge(table.rows[last_row_idx].cells[0])

            if is_military_day:
                top_left = table.rows[first_row_idx].cells[2]
                bottom_right = table.rows[last_row_idx].cells[3]
                top_left.merge(bottom_right)

                top_left.text = t['military']

                for paragraph in top_left.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(48)
                        run.font.name = 'Times New Roman'
                top_left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        doc.add_paragraph().add_run('\n')

        footer_table = doc.add_table(rows=1, cols=2)
        footer_table.autofit = True
        footer_table.width = section.page_width

        f_c1 = footer_table.cell(0, 0)
        fp1 = f_c1.paragraphs[0]
        fp1.add_run(t['director']).bold = True

        f_c2 = footer_table.cell(0, 1)
        fp2 = f_c2.paragraphs[0]
        fp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp2.add_run(director_name).bold = True

        f = BytesIO()
        doc.save(f)
        f.seek(0)

        filename = f"Jadval_{group.name}_{lang}.docx"
        response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response
    finally:
        translation.deactivate()



@user_passes_test(is_dean_or_admin)
def manage_plans(request):
    plans = AcademicPlan.objects.all().select_related('specialty', 'specialty__department__faculty', 'group')
    
    if hasattr(request.user, 'dean_profile'):
        faculty = request.user.dean_profile.faculty
        plans = plans.filter(
            Q(specialty__department__faculty=faculty) | 
            Q(group__specialty__department__faculty=faculty)
        )
        
    for plan in plans:
        if plan.group:
            plan.linked_groups = [plan.group]
        elif plan.specialty:
            plan.linked_groups = list(plan.specialty.groups.all())
        else:
            plan.linked_groups = []
        
    return render(request, 'schedule/plans/manage_plans.html', {'plans': plans})


@user_passes_test(is_dean_or_admin)
def create_plan(request):
    initial_data = {}
    specialty_id = request.GET.get('specialty')
    if specialty_id:
        initial_data['specialty'] = specialty_id

    if request.method == 'POST':
        form = AcademicPlanForm(request.POST, user=request.user)
        if form.is_valid():
            plan = form.save()
            return redirect('schedule:plan_detail', plan_id=plan.id)
    else:
        form = AcademicPlanForm(user=request.user, initial=initial_data)
    
    all_groups = Group.objects.all()
    if hasattr(request.user, 'dean_profile'):
        all_groups = all_groups.filter(specialty__department__faculty=request.user.dean_profile.faculty)

    return render(request, 'schedule/plans/create_plan.html', {
        'form': form, 
        'all_groups': all_groups
    })


@user_passes_test(is_dean_or_admin)
def plan_detail(request, plan_id):
    plan = get_object_or_404(AcademicPlan, id=plan_id)

    plan_faculty = None
    if plan.specialty and plan.specialty.department:
        plan_faculty = plan.specialty.department.faculty
    elif plan.group and plan.group.specialty and plan.group.specialty.department:
        plan_faculty = plan.group.specialty.department.faculty

    if hasattr(request.user, 'dean_profile'):
        user_faculty = request.user.dean_profile.faculty
        if plan_faculty and plan_faculty != user_faculty:
            messages.error(request, _("Это план чужого факультета!"))
            return redirect('schedule:manage_plans')
    else:
        user_faculty = plan_faculty

    try:
        current_sem_num = int(request.GET.get('semester', 1))
    except ValueError:
        current_sem_num = 1

    target_course = (current_sem_num + 1) // 2

    available_semesters = Semester.objects.filter(course=target_course)
    if user_faculty:
        available_semesters = available_semesters.filter(Q(faculty=user_faculty) | Q(faculty__isnull=True))

    available_semesters = available_semesters.order_by('-academic_year', 'number')
    active_semester = available_semesters.filter(is_active=True).first()

    credit_types = CreditType.objects.filter(Q(faculty=user_faculty) | Q(faculty__isnull=True))

    if request.method == 'POST' and 'add_discipline' in request.POST:
        form = PlanDisciplineForm(request.POST, faculty=user_faculty)
        if form.is_valid():
            disc = form.save(commit=False)
            disc.plan = plan
            disc.semester_number = current_sem_num
            disc.save()
            messages.success(request, _("Дисциплина успешно добавлена!"))
            return redirect(f"{request.path}?semester={current_sem_num}")
        else:
            messages.error(request, f"Ошибка при добавлении: проверьте правильность заполнения всех полей.")
    else:
        form = PlanDisciplineForm(initial={'semester_number': current_sem_num}, faculty=user_faculty)

    disciplines = PlanDiscipline.objects.filter(plan=plan, semester_number=current_sem_num)

    total_semester_credits = disciplines.aggregate(Sum('credits'))['credits__sum'] or 0
    total_semester_hours = disciplines.aggregate(
        total=Sum('lecture_hours') + Sum('practice_hours') + Sum('lab_hours') + Sum('control_hours') + Sum('independent_hours')
    )['total'] or 0

    credit_templates = CreditTemplate.objects.filter(Q(faculty=user_faculty) | Q(faculty__isnull=True)).order_by('credits')

    inst = plan_faculty.institute if plan_faculty else None

    return render(request, 'schedule/plans/plan_detail.html', {
        'plan': plan,
        'disciplines': disciplines,
        'form': form,
        'current_sem_num': current_sem_num,
        'semesters_range': range(1, 9),
        'target_course': target_course,
        'available_semesters': available_semesters,
        'active_semester': active_semester,
        'credit_types': credit_types,
        'credit_templates': credit_templates,
        'total_semester_credits': total_semester_credits,
        'total_semester_hours': total_semester_hours,
        'institute_pair_ratio': institute_pair_ratio(inst),
    })



@user_passes_test(is_dean_or_admin)
def generate_subjects_from_rup(request):
    teachers = Teacher.objects.select_related('user').all() 
    groups = Group.objects.all()
    if hasattr(request.user, 'dean_profile'):
        faculty = request.user.dean_profile.faculty
        groups = groups.filter(specialty__department__faculty=faculty)
        teachers = teachers.filter(
            Q(department__faculty=faculty) | 
            Q(additional_departments__faculty=faculty) |
            Q(subject__department__faculty=faculty) |
            Q(subject__groups__specialty__department__faculty=faculty) |
            Q(scheduleslot__group__specialty__department__faculty=faculty)
        ).distinct()

    suggestions = {}
    diagnostics =[]

    for group in groups:
        active_semester = get_active_semester_for_group(group)
        if not active_semester:
            diagnostics.append(f"Группа {group.name} ({group.course} курс): Нет активного семестра для этого курса.")
            continue

        current_semester_num = (group.course - 1) * 2 + active_semester.number

        plan = AcademicPlan.objects.filter(group=group, is_active=True).order_by('-admission_year').first()
        if not plan and group.specialty:
            plan = AcademicPlan.objects.filter(specialty=group.specialty, is_active=True).order_by('-admission_year').first()

        if not plan:
            diagnostics.append(f"Группа {group.name}: Нет активного РУП (учебного плана).")
            continue

        disciplines = PlanDiscipline.objects.filter(plan=plan, semester_number=current_semester_num)

        if not disciplines.exists():
            diagnostics.append(f"Группа {group.name}: В РУП '{plan.specialty.name if plan.specialty else plan.group.name}' нет дисциплин для {current_semester_num} семестра.")
            continue

        for disc in disciplines:
            exists = Subject.objects.filter(
                plan_discipline=disc,
                groups=group
            ).exists()

            if exists:
                continue

            key = f"{disc.subject_template.id}_{disc.id}_{active_semester.id}"

            if key not in suggestions:
                suggestions[key] = {
                    'key': key,
                    'template': disc.subject_template,
                    'discipline': disc,
                    'groups':[],
                    'hours': f"{disc.lecture_hours}/{disc.practice_hours}/{disc.lab_hours}/{disc.control_hours}",
                    'semester': active_semester
                }

            suggestions[key]['groups'].append(group)

    if request.method == 'POST':
        created_count = 0
        errors =[]
        selected_keys = request.POST.getlist('selected_items')

        if not selected_keys:
            errors.append("Вы ничего не выбрали или значения чекбоксов пустые (ошибка шаблона).")

        try:
            with transaction.atomic():
                for key in selected_keys:
                    if key not in suggestions:
                        errors.append(f"Ключ '{key}' не найден в списке. Возможно, страница устарела. Обновите страницу.")
                        continue

                    item = suggestions[key]
                    disc = item['discipline']
                    group_list = item['groups']
                    active_sem = item['semester']

                    is_stream = len(group_list) > 1 and request.POST.get(f'make_stream_{key}') == 'on'

                    teacher_id = request.POST.get(f'teacher_{key}')
                    teacher_obj = None
                    if teacher_id and str(teacher_id).isdigit():
                        teacher_obj = Teacher.objects.filter(id=int(teacher_id)).first()

                    target_dept = None
                    if group_list[0].specialty and group_list[0].specialty.department:
                        target_dept = group_list[0].specialty.department
                    elif hasattr(request.user, 'dean_profile') and request.user.dean_profile.faculty:
                        target_dept = Department.objects.filter(faculty=request.user.dean_profile.faculty).first()
                    else:
                        target_dept = Department.objects.first()

                    from journal.models import MatrixStructure
                    matrix = MatrixStructure.get_or_create_default(institute=target_dept.faculty.institute, faculty=None)
                    actual_weeks = matrix.columns.filter(col_type='WEEK').count() if matrix else 16

                    import uuid
                    if is_stream:
                        short_uuid = uuid.uuid4().hex[:6]
                        try:
                            subject = Subject.objects.create(
                                name=disc.subject_template.name,
                                code=f"STR{disc.id}-{short_uuid}",
                                department=target_dept,
                                type='LECTURE',
                                lecture_hours=disc.lecture_hours,
                                practice_hours=disc.practice_hours,
                                lab_hours=disc.lab_hours,
                                control_hours=disc.control_hours,
                                independent_work_hours=disc.independent_hours,
                                credits=disc.credits,
                                plan_discipline=disc,
                                is_stream_subject=True,
                                preferred_room_type=disc.preferred_room_type,
                                teacher=teacher_obj,
                                semester_weeks=actual_weeks
                            )
                            subject.groups.set(group_list)
                            created_count += 1
                        except Exception as e:
                            logger.exception("generate_subjects_from_plan stream")
                            errors.append(_("Ошибка БД при создании потока. См. журнал сервера."))
                    else:
                        for grp in group_list:
                            short_uuid = uuid.uuid4().hex[:6]

                            grp_dept = target_dept
                            if grp.specialty and grp.specialty.department:
                                grp_dept = grp.specialty.department

                            try:
                                subject = Subject.objects.create(
                                    name=disc.subject_template.name,
                                    code=f"S{disc.id}-{short_uuid}",
                                    department=grp_dept,
                                    type='LECTURE',
                                    lecture_hours=disc.lecture_hours,
                                    practice_hours=disc.practice_hours,
                                    lab_hours=disc.lab_hours,
                                    control_hours=disc.control_hours,
                                    independent_work_hours=disc.independent_hours,
                                    credits=disc.credits,
                                    plan_discipline=disc,
                                    is_stream_subject=False,
                                    preferred_room_type=disc.preferred_room_type,
                                    teacher=teacher_obj,
                                    semester_weeks=actual_weeks
                                )
                                subject.groups.add(grp)
                                created_count += 1
                            except Exception as e:
                                logger.exception("generate_subjects_from_plan group %s", grp.name)
                                errors.append(_("Ошибка БД при создании предмета для группы %(name)s. См. журнал сервера.") % {'name': grp.name})

        except Exception as e:
            logger.exception("generate_subjects_from_plan")
            errors.append(_("Критическая ошибка. См. журнал сервера."))

        if errors:
            for err in errors:
                messages.error(request, err)

        if created_count > 0:
            messages.success(request, f"Успешно создано предметов: {created_count}")
            return redirect('schedule:manage_subjects')
        else:
            messages.warning(request, "Предметы не были созданы. Посмотрите красные ошибки выше.")

    return render(request, 'schedule/plans/generate_preview.html', {
        'suggestions': suggestions.values(),
        'diagnostics': diagnostics,
        'teachers': teachers,
    })



@user_passes_test(is_facility_admin)
def edit_classroom(request, classroom_id):
    classroom = get_object_or_404(Classroom, id=classroom_id)
    if request.method == 'POST':
        form = ClassroomForm(request.POST, instance=classroom)
        if form.is_valid():
            form.save()
            messages.success(request, _('Кабинет %(classroom_number)s обновлен') % {'classroom_number': classroom.number})
            return redirect('schedule:manage_classrooms')
    else:
        form = ClassroomForm(instance=classroom)
    return render(request, 'schedule/classroom_form.html', {
        'form': form, 
        'title': _('Редактировать кабинет %(classroom_number)s') % {'classroom_number': classroom.number}
    })

@login_required
@login_required
def classroom_occupancy(request):
    day = int(request.GET.get('day', 0))
    shift_filter = request.GET.get('shift', '')  
 
    active_semesters = Semester.objects.filter(is_active=True)
    if not active_semesters.exists():
        active_semesters = Semester.objects.order_by('-start_date')[:1]
 
    classrooms = Classroom.objects.select_related('building').filter(is_active=True)
    institutes = []
    selected_institute_id = request.GET.get('institute')
 
    if request.user.is_superuser or hasattr(request.user, 'director_profile') or hasattr(request.user, 'prorector_profile'):
        institutes = Institute.objects.all()
        if selected_institute_id:
            classrooms = classrooms.filter(building__institute_id=selected_institute_id)
    elif hasattr(request.user, 'dean_profile'):
        faculty = request.user.dean_profile.faculty
        if faculty and faculty.institute:
            classrooms = classrooms.filter(building__institute=faculty.institute)
            institutes = []
 
    classrooms = classrooms.order_by('building', 'floor', 'number')
 
    institute_for_ts = None
    if selected_institute_id:
        institute_for_ts = Institute.objects.filter(id=selected_institute_id).first()
    elif hasattr(request.user, 'dean_profile') and request.user.dean_profile.faculty:
        institute_for_ts = request.user.dean_profile.faculty.institute
 
    if institute_for_ts:
        inst_ts = TimeSlot.objects.filter(institute=institute_for_ts).order_by('start_time')
        time_slots = inst_ts if inst_ts.exists() else TimeSlot.objects.filter(institute__isnull=True).order_by('start_time')
    else:
        time_slots = TimeSlot.objects.filter(institute__isnull=True).order_by('start_time')
 
    if shift_filter in ('MORNING', 'DAY', 'EVENING'):
        time_slots = time_slots.filter(shift=shift_filter)
 
    days = [
        (0, _('Понедельник')), (1, _('Вторник')), (2, _('Среда')),
        (3, _('Четверг')), (4, _('Пятница')), (5, _('Суббота'))
    ]
 
    slots = ScheduleSlot.objects.filter(
        semester__in=active_semesters,
        day_of_week=day,
        is_active=True,
        classroom__in=classrooms
    ).select_related('classroom', 'group', 'subject', 'time_slot', 'teacher__user')
 
    occupancy = {}
    for slot in slots:
        c_id = slot.classroom.id
        t_id = slot.time_slot.id
        if c_id not in occupancy:
            occupancy[c_id] = {}
        if t_id not in occupancy[c_id]:
            occupancy[c_id][t_id] = slot
 
    return render(request, 'schedule/classroom_occupancy.html', {
        'classrooms': classrooms,
        'time_slots': time_slots,
        'occupancy': occupancy,
        'selected_day': day,
        'days': days,
        'institutes': institutes,
        'selected_institute_id': int(selected_institute_id) if selected_institute_id else None,
        'shift_filter': shift_filter,
    })


@login_required
def import_schedule_view(request):
    if not (hasattr(request.user, 'dean_profile') or hasattr(request.user, 'vicedean_profile') or request.user.is_superuser):
        messages.error(request, _("Доступ запрещен"))
        return redirect('schedule:view')

    if request.method == 'POST' and 'confirm_import' in request.POST:
        semester_id = request.POST.get('semester_id')
        semester = Semester.objects.get(id=semester_id)

        all_keys = [k for k in request.POST.keys() if k.startswith('item_') and k.endswith('_group_id')]
        if not all_keys:
            messages.error(request, _("Нет данных для сохранения"))
            return redirect('schedule:view')

        import_data = []

        for key in all_keys:
            index = key.split('_')[1]
            prefix = f'item_{index}_'

            group_id = request.POST.get(f'{prefix}group_id')
            if not group_id: continue  
            import_data.append({
                'group_id': group_id,
                'day': request.POST.get(f'{prefix}day'),
                'start_time': request.POST.get(f'{prefix}start_time'),
                'end_time': request.POST.get(f'{prefix}end_time'),
                'subject': request.POST.get(f'{prefix}subject'),
                'teacher': request.POST.get(f'{prefix}teacher'),
                'room': request.POST.get(f'{prefix}room'),
                'type': request.POST.get(f'{prefix}type'),
                'is_military': request.POST.get(f'{prefix}is_military'),
            })

        importer = ScheduleImporter()
        stats = importer.save_data_from_preview(import_data, semester)
        messages.success(request, _("✅ Импорт завершен! Слотов: %(slots)s, Новых предметов: %(subjects)s, Новых учителей: %(teachers)s") % {
            'slots': stats['slots'],
            'subjects': stats['subjects'],
            'teachers': stats['teachers']
        })
        return redirect('schedule:view')

    elif request.method == 'POST':
        form = ScheduleImportForm(request.POST, request.FILES)
        if form.is_valid():
            try:
                importer = ScheduleImporter(file=request.FILES['file'])
                preview_data = importer.parse_for_preview(default_group=form.cleaned_data['default_group'])

                if not preview_data:
                    messages.warning(request, _("Данные не найдены в файле или формат некорректен."))
                    return redirect('schedule:import_schedule')

                context = {
                    'preview_data': preview_data,
                    'semester': form.cleaned_data['semester'],
                    'days_map_rev': {0: _('Понедельник'), 1: _('Вторник'), 2: _('Среда'), 3: _('Четверг'), 4: _('Пятница'), 5: _('Суббота')}
                }
                return render(request, 'schedule/import_preview_edit.html', context)

            except Exception as e:
                logger.exception("import_schedule parse")
                messages.error(request, _("Ошибка обработки файла. Попробуйте ещё раз или обратитесь к администратору."))

    else:
        form = ScheduleImportForm()

    return render(request, 'schedule/import_schedule.html', {'form': form})

@login_required
@require_POST
def api_create_subject_template(request):
    
    import json
    try:
        data = json.loads(request.body)
        name = data.get('name', '').strip()
        
        if not name:
            return JsonResponse({'success': False, 'error': _('Название не может быть пустым')})
            
        if SubjectTemplate.objects.filter(name__iexact=name).exists():
            return JsonResponse({'success': False, 'error': _('Такая дисциплина уже существует')})
            
        template = SubjectTemplate.objects.create(name=name)
        
        return JsonResponse({
            'success': True,
            'id': template.id,
            'name': template.name
        })
    except Exception as e:
        logger.exception("api_create_subject_template")
        return JsonResponse({'success': False, 'error': _('Внутренняя ошибка сервера')})


@login_required
@user_passes_test(is_dean_or_admin)
def copy_plan(request, plan_id):
    original_plan = get_object_or_404(AcademicPlan, id=plan_id)
    
    if request.method == 'POST':
        new_year = request.POST.get('new_year')
        if not new_year:
            messages.error(request, _("Укажите новый год"))
            return redirect('schedule:manage_plans')
            
        if original_plan.specialty:
            exists = AcademicPlan.objects.filter(specialty=original_plan.specialty, admission_year=new_year).exists()
            plan_name = original_plan.specialty.code
        else:
            exists = AcademicPlan.objects.filter(group=original_plan.group, admission_year=new_year).exists()
            plan_name = original_plan.group.name

        if exists:
            messages.error(request, _("План для %(plan_name)s на %(new_year)s год уже существует!") % {
                'plan_name': plan_name,
                'new_year': new_year
            })
            return redirect('schedule:manage_plans')

        with transaction.atomic():
            new_plan = AcademicPlan.objects.create(
                specialty=original_plan.specialty,
                group=original_plan.group,  
                admission_year=new_year,
                is_active=True
            )
            
            disciplines = original_plan.disciplines.all()
            new_disciplines =[]
            for disc in disciplines:
                new_disciplines.append(PlanDiscipline(
                    plan=new_plan,
                    subject_template=disc.subject_template,
                    semester_number=disc.semester_number,
                    cycle=disc.cycle,
                    has_subgroups=disc.has_subgroups,
                    discipline_type=disc.discipline_type,
                    credits=disc.credits,
                    lecture_hours=disc.lecture_hours,
                    practice_hours=disc.practice_hours,
                    lab_hours=disc.lab_hours,
                    control_hours=disc.control_hours,
                    independent_hours=disc.independent_hours,
                    control_type=disc.control_type,
                    has_course_work=disc.has_course_work,
                    preferred_room_type=disc.preferred_room_type
                ))
            PlanDiscipline.objects.bulk_create(new_disciplines)
            
        messages.success(request, _("План успешно скопирован на %(new_year)s год!") % {'new_year': new_year})
        return redirect('schedule:plan_detail', plan_id=new_plan.id)
        
    return redirect('schedule:manage_plans')





@user_passes_test(is_dean_or_admin)
def delete_plan_discipline(request, discipline_id):
    discipline = get_object_or_404(PlanDiscipline, id=discipline_id)
    plan_id = discipline.plan.id
    
    if hasattr(request.user, 'dean_profile'):
        plan = discipline.plan
        plan_faculty = None
        if plan.specialty and plan.specialty.department:
            plan_faculty = plan.specialty.department.faculty
        elif plan.group and plan.group.specialty and plan.group.specialty.department:
            plan_faculty = plan.group.specialty.department.faculty
            
        if plan_faculty and plan_faculty != request.user.dean_profile.faculty:
            messages.error(request, _("Нет прав на удаление"))
            return redirect('schedule:plan_detail', plan_id=plan_id)

    discipline.delete()
    messages.success(request, _("Дисциплина удалена из плана"))
    return redirect('schedule:plan_detail', plan_id=plan_id)

@user_passes_test(is_dean_or_admin)
def teacher_load_report(request):
    teachers = Teacher.objects.select_related('user', 'department', 'department__faculty').prefetch_related('subject_set').all()
    
    if hasattr(request.user, 'dean_profile'):
        faculty = request.user.dean_profile.faculty
        teachers = teachers.filter(department__faculty=faculty)
    
    report_data = []
    
    for teacher in teachers:
        subjects = teacher.subject_set.all()
        
        total_hours = 0
        total_credits = 0
        subjects_list = []
        
        for subj in subjects:
            hours = subj.total_auditory_hours
            total_hours += hours
            total_credits += subj.credits
            subjects_list.append(f"{subj.name} ({hours}ч)")
            
        report_data.append({
            'teacher': teacher,
            'total_hours': total_hours,
            'total_credits': total_credits,
            'subjects_count': subjects.count(),
            'subjects_names': ", ".join(subjects_list[:3]) + ("..." if len(subjects_list) > 3 else "")
        })
    
    report_data.sort(key=lambda x: x['total_hours'], reverse=True)
    
    return render(request, 'schedule/teacher_load_report.html', {
        'report_data': report_data
    })

@login_required
def subject_materials(request, subject_id):
    subject = get_object_or_404(Subject, id=subject_id)
    materials = subject.materials.order_by('-uploaded_at')
    
    can_upload = False
    if request.user.is_superuser or request.user.role in ['DEAN', 'VICE_DEAN']:
        can_upload = True
    elif hasattr(request.user, 'teacher_profile') and subject.teacher and subject.teacher.user == request.user:
        can_upload = True

    if request.method == 'POST' and can_upload:
        form = MaterialUploadForm(request.POST, request.FILES)
        if form.is_valid():
            mat = form.save(commit=False)
            mat.subject = subject
            mat.save()
            messages.success(request, _('Материал добавлен'))
            return redirect('schedule:subject_materials', subject_id=subject.id)
    else:
        form = MaterialUploadForm()

    return render(request, 'schedule/subject_materials.html', {
        'subject': subject,
        'materials': materials,
        'form': form,
        'can_upload': can_upload
    })

@login_required
def delete_material(request, material_id):
    material = get_object_or_404(SubjectMaterial, id=material_id)
    if request.user.is_superuser or (hasattr(request.user, 'teacher_profile') and material.subject.teacher.user == request.user):
        material.delete()
        messages.success(request, _('Материал удален'))
    return redirect('schedule:subject_materials', subject_id=material.subject.id)



@user_passes_test(is_dean_or_admin)
def activate_plan(request, plan_id):
    plan = get_object_or_404(AcademicPlan, id=plan_id)
    
    with transaction.atomic():
        if plan.specialty:
            AcademicPlan.objects.filter(specialty=plan.specialty, group__isnull=True).update(is_active=False)
        elif plan.group:
            AcademicPlan.objects.filter(group=plan.group).update(is_active=False)
        
        plan.is_active = True
        plan.save()
        
    if plan.specialty:
        messages.success(request, _("Учебный план для %(specialty_code)s (%(admission_year)s) теперь АКТИВЕН!") % {
            'specialty_code': plan.specialty.code,
            'admission_year': plan.admission_year
        })
    else:
        messages.success(request, _("Учебный план для группы %(group_name)s (%(admission_year)s) теперь АКТИВЕН!") % {
            'group_name': plan.group.name,
            'admission_year': plan.admission_year
        })
    return redirect('schedule:plan_detail', plan_id=plan.id)

@user_passes_test(is_dean_or_admin)
def delete_plan(request, plan_id):
    plan = get_object_or_404(AcademicPlan, id=plan_id)
    if hasattr(request.user, 'dean_profile'):
        plan_faculty = None
        if plan.specialty and plan.specialty.department:
            plan_faculty = plan.specialty.department.faculty
        elif plan.group and plan.group.specialty and plan.group.specialty.department:
            plan_faculty = plan.group.specialty.department.faculty
            
        if plan_faculty and plan_faculty != request.user.dean_profile.faculty:
            messages.error(request, _("Нет прав на удаление этого плана"))
            return redirect('schedule:manage_plans')
            
    plan.delete()
    messages.success(request, _("Учебный план удален"))
    return redirect('schedule:manage_plans')

@user_passes_test(is_dean_or_admin)
def activate_semester_from_plan(request, semester_id):
    semester = get_object_or_404(Semester, id=semester_id)
    
    with transaction.atomic():
        Semester.objects.filter(course=semester.course, is_active=True).update(is_active=False)
        
        semester.is_active = True
        semester.save()
        
    messages.success(request, _("Семестр '%(semester_name)s' (%(course)s курс) успешно активирован!") % {
        'semester_name': semester.name,
        'course': semester.course
    })
    return redirect(request.META.get('HTTP_REFERER', 'schedule:manage_plans'))

@user_passes_test(is_dean_or_admin)
@require_POST
def set_active_semester_manual(request):
    semester_id = request.POST.get('semester_id')
    
    if not semester_id:
        messages.error(request, _("Выберите семестр из списка"))
        return redirect(request.META.get('HTTP_REFERER'))

    semester = get_object_or_404(Semester, id=semester_id)
    
    if hasattr(request.user, 'dean_profile') and semester.faculty != request.user.dean_profile.faculty:
        messages.error(request, _("Ошибка доступа"))
        return redirect(request.META.get('HTTP_REFERER'))

    with transaction.atomic():
        Semester.objects.filter(
            faculty=semester.faculty, 
            course=semester.course, 
            is_active=True
        ).update(is_active=False)
        
        semester.is_active = True
        semester.save()

    messages.success(request, _("Семестр %(semester)s теперь АКТИВЕН для %(course)s курса!") % {
        'semester': semester,
        'course': semester.course
    })
    return redirect(request.META.get('HTTP_REFERER'))


@user_passes_test(lambda u: u.is_superuser or hasattr(u, 'director_profile') or hasattr(u, 'prorector_profile'))
def manage_time_slots(request):
    if request.method == 'POST':
        form = TimeSlotGeneratorForm(request.POST)
        if form.is_valid():
            data = form.cleaned_data
            inst = data['institute'] 
            shift = data['shift']
            
            if data['delete_existing']:
                TimeSlot.objects.filter(institute=inst, shift=shift).delete()
            
            current_time = datetime.combine(date.today(), data['start_time'])
            created_count = 0
            
            big_break_after = data.get('big_break_after') or 0
            big_break_dur = data.get('big_break_duration') or data['break_duration']
            
            for i in range(1, data['pairs_count'] + 1):
                start = current_time.time()
                end_dt = current_time + timedelta(minutes=data['lesson_duration'])
                end = end_dt.time()
                
                TimeSlot.objects.create(
                    institute=inst,
                    number=i,
                    start_time=start,
                    end_time=end,
                    shift=shift,
                    duration=data['lesson_duration']
                )
                created_count += 1
                
                is_big_break = (i == big_break_after)
                break_min = big_break_dur if is_big_break else data['break_duration']
                
                current_time = end_dt + timedelta(minutes=break_min)
            
            target = inst.name if inst else "Всего Университета (Глобально)"
            messages.success(request, _(f"Сетка звонков создана для: {target}. Сгенерировано {created_count} пар."))
            return redirect('schedule:manage_time_slots')
    else:
        form = TimeSlotGeneratorForm()

    global_slots = TimeSlot.objects.filter(institute__isnull=True).order_by('shift', 'start_time')
    institute_slots = TimeSlot.objects.filter(institute__isnull=False).select_related('institute').order_by('institute', 'shift', 'start_time')
    
    return render(request, 'schedule/manage_time_slots.html', {
        'form': form,
        'global_slots': global_slots,
        'institute_slots': institute_slots
    })




@user_passes_test(is_facility_admin)
def manage_buildings(request):
    buildings = Building.objects.select_related('institute').all()
    
    if hasattr(request.user, 'dean_profile'):
        faculty = request.user.dean_profile.faculty
        if faculty and faculty.institute:
            buildings = buildings.filter(institute=faculty.institute)

    return render(request, 'schedule/manage_buildings.html', {'buildings': buildings})

@user_passes_test(is_facility_admin)
def add_building(request):
    if request.method == 'POST':
        form = BuildingForm(request.POST, user=request.user)
        if form.is_valid():
            form.save()
            messages.success(request, _("Учебный корпус добавлен"))
            return redirect('schedule:manage_buildings')
    else:
        form = BuildingForm(user=request.user)
    
    return render(request, 'core/form_generic.html', {
        'form': form, 
        'title': _('Добавить учебный корпус'),
        'cancel_url': reverse('schedule:manage_buildings')
    })

@user_passes_test(is_facility_admin)
def edit_building(request, building_id):
    building = get_object_or_404(Building, id=building_id)
    
    if hasattr(request.user, 'dean_profile'):
        faculty = request.user.dean_profile.faculty
        if building.institute != faculty.institute:
            messages.error(request, _("Вы не можете редактировать корпуса чужого института"))
            return redirect('schedule:manage_buildings')

    if request.method == 'POST':
        form = BuildingForm(request.POST, instance=building, user=request.user)
        if form.is_valid():
            form.save()
            messages.success(request, _("Корпус обновлен"))
            return redirect('schedule:manage_buildings')
    else:
        form = BuildingForm(instance=building, user=request.user)
    
    return render(request, 'core/form_generic.html', {
        'form': form, 
        'title': _('Редактировать корпус'),
        'cancel_url': reverse('schedule:manage_buildings')
    })

@user_passes_test(is_facility_admin)
def delete_building(request, building_id):
    building = get_object_or_404(Building, id=building_id)
    
    if hasattr(request.user, 'dean_profile'):
        faculty = request.user.dean_profile.faculty
        if building.institute != faculty.institute:
            messages.error(request, _("Нет доступа"))
            return redirect('schedule:manage_buildings')
            
    if building.classrooms.exists():
        messages.error(request, _("Нельзя удалить корпус, в котором есть кабинеты. Сначала удалите кабинеты."))
    else:
        building.delete()
        messages.success(request, _("Корпус удален"))
        
    return redirect('schedule:manage_buildings')



@login_required
@require_POST
def check_schedule_conflicts(request):
    import json
    data = json.loads(request.body)
    
    semester_id = data.get('semester_id')
    day = data.get('day')
    time_slot_id = data.get('time_slot_id')
    group_id = data.get('group_id')
    teacher_id = data.get('teacher_id') 
    room_number = data.get('room') 
    subject_id = data.get('subject_id')
    
    conflicts = []
    
    base_qs = ScheduleSlot.objects.filter(
        semester_id=semester_id,
        day_of_week=day,
        time_slot_id=time_slot_id,
        is_active=True
    )

    if group_id:
        group_conflict = base_qs.filter(group_id=group_id).first()
        if group_conflict:
            conflicts.append({
                'type': 'group',
                'message': _(f"Группа {group_conflict.group.name} уже занята: {group_conflict.subject.name}")
            })

    if teacher_id:
        teacher_conflict = base_qs.filter(teacher_id=teacher_id).first()
        if teacher_conflict:
            conflicts.append({
                'type': 'teacher',
                'message': _(f"Преподаватель {teacher_conflict.teacher.user.get_full_name()} уже ведет пару в группе {teacher_conflict.group.name}")
            })

    if room_number:
        
        classroom_qs = Classroom.objects.filter(number=room_number, is_active=True)
        
        if group_id:
            try:
                grp = Group.objects.get(id=group_id)
                inst = grp.specialty.department.faculty.institute
                classroom_qs = classroom_qs.filter(building__institute=inst)
            except Exception:
                logger.exception("check_conflicts classroom filter")
        
        target_classroom = classroom_qs.first()
        
        if target_classroom:
            room_conflict = base_qs.filter(classroom=target_classroom).first()
            if room_conflict:
                conflicts.append({
                    'type': 'room',
                    'message': _(f"Кабинет {room_number} ({target_classroom.building.name}) занят: {room_conflict.group.name}, {room_conflict.teacher.user.last_name if room_conflict.teacher else ''}")
                })
        else:
            text_conflict = base_qs.filter(room=room_number).first()
            if text_conflict:
                 conflicts.append({
                    'type': 'room',
                    'message': _(f"Кабинет {room_number} (текст) занят: {text_conflict.group.name}")
                })

    if conflicts:
        return JsonResponse({'success': False, 'conflicts': conflicts})
    
    return JsonResponse({'success': True})

@user_passes_test(lambda u: u.is_superuser or u.role in [
    'DEAN', 'VICE_DEAN', 'HEAD_OF_DEPT', 'DIRECTOR', 'PRO_RECTOR'
])
@user_passes_test(lambda u: u.is_superuser or u.role in [
    'DEAN', 'VICE_DEAN', 'HEAD_OF_DEPT', 'DIRECTOR', 'PRO_RECTOR'
])
def import_rup_excel(request, plan_id, semester_num):
    from .services import RupImporter
 
    plan = get_object_or_404(AcademicPlan, id=plan_id)
    plan_url = f"/schedule/plans/{plan.id}/?semester={semester_num}"
 
    if plan.group:
        plan_groups = Group.objects.filter(id=plan.group.id)
    elif plan.specialty:
        plan_groups = Group.objects.filter(specialty=plan.specialty).order_by('course', 'name')
    else:
        plan_groups = Group.objects.none()
 
    teachers = _get_teachers(request.user, plan)
 
    if request.method == 'POST' and 'confirm_import' in request.POST:
        preview_data = request.session.get('rup_import_preview', [])
        if not preview_data:
            messages.error(request, "Сессия истекла. Загрузите файл заново.")
            return redirect(plan_url)
 
        stats = {'disciplines': 0, 'subjects': 0, 'errors': []}
 
        try:
            with transaction.atomic():
                for item in preview_data:
                    idx = str(item['id'])
 
                    if not request.POST.get(f"import_{idx}"):
                        continue
 
                    name = request.POST.get(f"name_{idx}", item.get('name', '')).strip()
                    if not name:
                        continue
 
                    group_ids = request.POST.getlist(f"groups_{idx}")
                    if not group_ids:
                        stats['errors'].append(f"«{name}» — не выбрана группа, пропущено.")
                        continue
 
                    disc_type = request.POST.get(f"type_{idx}", item.get('type', 'REQUIRED'))
                    credits   = safe_int(request.POST.get(f"credits_{idx}", item.get('credits', 0)))
                    lec       = safe_int(request.POST.get(f"lec_{idx}",     item.get('lec', 0)))
                    prac      = safe_int(request.POST.get(f"prac_{idx}",    item.get('prac', 0)))
                    srsp      = safe_int(request.POST.get(f"srsp_{idx}",    item.get('srsp', 0)))
                    srs       = safe_int(request.POST.get(f"srs_{idx}",     item.get('srs', 0)))
                    sem_num   = safe_int(request.POST.get(f"semester_{idx}", semester_num)) or semester_num
 
                    teacher1_id = request.POST.get(f"teacher1_{idx}", "")
                    teacher1 = Teacher.objects.filter(id=teacher1_id).first() if teacher1_id else None
 
                    do_split = bool(request.POST.get(f"split_{idx}"))
                    teacher2 = None
                    split_lec = split_prac = split_srsp = split_srs = split_credits = 0
 
                    if do_split:
                        teacher2_id = request.POST.get(f"teacher2_{idx}", "")
                        teacher2 = Teacher.objects.filter(id=teacher2_id).first() if teacher2_id else None
                        split_lec     = safe_int(request.POST.get(f"split_lec_{idx}",     0))
                        split_prac    = safe_int(request.POST.get(f"split_prac_{idx}",    0))
                        split_srsp    = safe_int(request.POST.get(f"split_srsp_{idx}",    0))
                        split_srs     = safe_int(request.POST.get(f"split_srs_{idx}",     0))
                        split_credits = safe_int(request.POST.get(f"split_credits_{idx}", 0))
 
                        if not teacher2:
                            stats['errors'].append(
                                f"«{name}» — разделение включено, но второй преподаватель не выбран. "
                                f"Дисциплина создана без разделения."
                            )
                            do_split = False
 
                        over_lec  = split_lec  > lec
                        over_prac = split_prac > prac
                        over_srsp = split_srsp > srsp
                        over_srs  = split_srs  > srs
                        if over_lec or over_prac or over_srsp or over_srs:
                            stats['errors'].append(
                                f"«{name}» — часы для второго преподавателя превышают общие. "
                                f"Разделение отменено, дисциплина создана без разделения."
                            )
                            do_split = False
 
                    try:
                        template, _ = SubjectTemplate.objects.get_or_create(name=name)
                        discipline, _ = PlanDiscipline.objects.update_or_create(
                            plan=plan,
                            subject_template=template,
                            semester_number=sem_num,
                            defaults={
                                'discipline_type':   disc_type,
                                'credits':           credits,
                                'lecture_hours':     lec,
                                'practice_hours':    prac,
                                'lab_hours':         0,
                                'control_hours':     srsp,
                                'independent_hours': srs,
                                'control_type':      'EXAM',
                                'cycle':             'OTHER',
                            }
                        )
                        stats['disciplines'] += 1
 
                        selected_groups = Group.objects.filter(id__in=group_ids)

                        for grp in selected_groups:
                            already = Subject.objects.filter(
                                plan_discipline=discipline, groups=grp
                            ).exists()
                            if already:
                                continue
 
                            dept = _resolve_dept(grp, plan)
                            if not dept:
                                stats['errors'].append(
                                    f"«{name}» / {grp.name} — кафедра не найдена."
                                )
                                continue

                            from journal.models import MatrixStructure
                            matrix = MatrixStructure.get_or_create_default(institute=dept.faculty.institute, faculty=None)
                            actual_weeks = matrix.columns.filter(col_type='WEEK').count() if matrix else 16
 
                            if do_split:
                                main_lec  = lec  - split_lec
                                main_prac = prac - split_prac
                                main_srsp = srsp - split_srsp
                                main_srs  = srs  - split_srs
                                main_cred = max(0, credits - split_credits)
 
                                subj1 = Subject.objects.create(
                                    name=name,
                                    code=f"RUP{discipline.id}-{grp.id}-{uuid.uuid4().hex[:4]}",
                                    department=dept,
                                    type='LECTURE',
                                    lecture_hours=main_lec,
                                    practice_hours=main_prac,
                                    lab_hours=0,
                                    control_hours=main_srsp,
                                    independent_work_hours=main_srs,
                                    credits=main_cred,
                                    plan_discipline=discipline,
                                    preferred_room_type=discipline.preferred_room_type,
                                    teacher=teacher1,
                                    is_active=True,
                                    semester_weeks=actual_weeks,
                                )
                                subj1.groups.add(grp)
 
                                subj2 = Subject.objects.create(
                                    name=f"{name} (ч.2)",
                                    code=f"RUP{discipline.id}-{grp.id}-{uuid.uuid4().hex[:4]}-2",
                                    department=dept,
                                    type='PRACTICE',
                                    lecture_hours=split_lec,
                                    practice_hours=split_prac,
                                    lab_hours=0,
                                    control_hours=split_srsp,
                                    independent_work_hours=split_srs,
                                    credits=split_credits,
                                    plan_discipline=discipline,
                                    preferred_room_type=discipline.preferred_room_type,
                                    teacher=teacher2,
                                    is_active=True,
                                    semester_weeks=actual_weeks,
                                )
                                subj2.groups.add(grp)
                                stats['subjects'] += 2
                            else:
                                subj = Subject.objects.create(
                                    name=name,
                                    code=f"RUP{discipline.id}-{grp.id}-{uuid.uuid4().hex[:4]}",
                                    department=dept,
                                    type='LECTURE',
                                    lecture_hours=lec,
                                    practice_hours=prac,
                                    lab_hours=0,
                                    control_hours=srsp,
                                    independent_work_hours=srs,
                                    credits=credits,
                                    plan_discipline=discipline,
                                    preferred_room_type=discipline.preferred_room_type,
                                    teacher=teacher1,
                                    is_active=True,
                                    semester_weeks=actual_weeks,
                                )
                                subj.groups.add(grp)
                                stats['subjects'] += 1
 
                    except Exception as e:
                        stats['errors'].append(f"«{name}»: {e}")
 
        except Exception as e:
            messages.error(request, f"Критическая ошибка: {e}")
            return redirect(plan_url)
        finally:
            request.session.pop('rup_import_preview', None)
 
        messages.success(
            request,
            f"Готово! Дисциплин в план: {stats['disciplines']}, "
            f"предметов создано: {stats['subjects']}."
        )
        if stats['errors']:
            messages.warning(
                request,
                f"Предупреждения ({len(stats['errors'])}): "
                + "; ".join(stats['errors'][:5])
                + ("..." if len(stats['errors']) > 5 else "")
            )
        return redirect(plan_url)
 
    if request.method == 'POST' and 'file' in request.FILES:
        try:
            preview_data = RupImporter.parse_for_preview(request.FILES['file'])
        except Exception as e:
            messages.error(request, f"Ошибка обработки файла: {e}")
            return redirect(plan_url)
 
        if not preview_data:
            messages.warning(request, "Файл не содержит распознанных дисциплин.")
            return redirect(plan_url)
 
        request.session['rup_import_preview'] = preview_data
        request.session.modified = True
 
        return render(request, 'schedule/plans/import_rup_preview.html', {
            'plan':         plan,
            'semester_num': semester_num,
            'preview_data': preview_data,
            'plan_groups':  plan_groups,
            'teachers':     teachers,
        })
 
    messages.error(request, "Неверный запрос.")
    return redirect(plan_url)

def _resolve_dept(grp, plan):
    if grp.specialty and grp.specialty.department:
        return grp.specialty.department
    if plan.specialty and plan.specialty.department:
        return plan.specialty.department
    return Department.objects.first()
 
 
def _get_teachers(user, plan):
    qs = Teacher.objects.select_related('user', 'department').order_by(
        'user__last_name', 'user__first_name'
    )
    faculty = None
    if hasattr(user, 'dean_profile') and user.dean_profile.faculty:
        faculty = user.dean_profile.faculty
    elif hasattr(user, 'vicedean_profile') and user.vicedean_profile.faculty:
        faculty = user.vicedean_profile.faculty
    elif hasattr(user, 'head_of_dept_profile') and user.head_of_dept_profile.department:
        return qs.filter(
            Q(department=user.head_of_dept_profile.department) |
            Q(additional_departments=user.head_of_dept_profile.department)
        ).distinct()
 
    if faculty:
        return qs.filter(
            Q(department__faculty=faculty) |
            Q(additional_departments__faculty=faculty)
        ).distinct()
 
    if plan.specialty and plan.specialty.department:
        dept = plan.specialty.department
        return qs.filter(
            Q(department=dept) |
            Q(department__faculty=dept.faculty) |
            Q(additional_departments=dept)
        ).distinct()
 
    return qs

@login_required
@user_passes_test(is_dean_or_admin)
@require_POST
def api_ai_assign_teachers(request):
    try:
        data = json.loads(request.body)
        subjects = data.get('subjects', [])
        teacher_ids = data.get('teacher_ids',[])
        model_name = data.get('model_name', 'gemma3:4b')
        assign_method = data.get('method', 'algo')
        
        if not subjects:
            return JsonResponse({'success': False, 'error': 'Нет предметов для распределения'})

        teachers_qs = Teacher.objects.select_related('user', 'department').prefetch_related('competencies', 'additional_departments').all()
        
        if hasattr(request.user, 'dean_profile') or hasattr(request.user, 'vicedean_profile'):
            profile = getattr(request.user, 'dean_profile', None) or getattr(request.user, 'vicedean_profile', None)
            faculty = profile.faculty
            institute = faculty.institute if faculty else None
            if institute:
                teachers_qs = teachers_qs.filter(
                    Q(department__faculty__institute=institute) | 
                    Q(additional_departments__faculty__institute=institute) |
                    Q(subject__department__faculty__institute=institute) |
                    Q(subject__groups__specialty__department__faculty__institute=institute) |
                    Q(scheduleslot__group__specialty__department__faculty__institute=institute)
                ).distinct()
            else:
                teachers_qs = teachers_qs.filter(
                    Q(department__faculty=faculty) | 
                    Q(additional_departments__faculty=faculty) |
                    Q(subject__department__faculty=faculty) |
                    Q(subject__groups__specialty__department__faculty=faculty) |
                    Q(scheduleslot__group__specialty__department__faculty=faculty)
                ).distinct()
        elif hasattr(request.user, 'director_profile') or hasattr(request.user, 'prorector_profile'):
            profile = getattr(request.user, 'director_profile', None) or getattr(request.user, 'prorector_profile', None)
            institute = profile.institute
            if institute:
                teachers_qs = teachers_qs.filter(
                    Q(department__faculty__institute=institute) | 
                    Q(additional_departments__faculty__institute=institute) |
                    Q(subject__department__faculty__institute=institute) |
                    Q(subject__groups__specialty__department__faculty__institute=institute) |
                    Q(scheduleslot__group__specialty__department__faculty__institute=institute)
                ).distinct()
            
        if teacher_ids:
            teachers_qs = teachers_qs.filter(id__in=teacher_ids)
            
        if assign_method == 'algo':
            result = AlgorithmicAssignmentService.generate_assignment(teachers_qs, subjects)
        else:
            result = AIAssignmentService.generate_assignment(teachers_qs, subjects, model_name)
        
        if result and "assignments" in result:
            return JsonResponse({'success': True, 'assignments': result["assignments"]})
        else:
            return JsonResponse({'success': False, 'error': 'Система не смогла сформировать корректный ответ.'})
            
    except Exception as e:
        logger.exception("api_assign_teachers")
        return JsonResponse({'success': False, 'error': _('Внутренняя ошибка сервера')})



@user_passes_test(lambda u: u.is_superuser or hasattr(u, 'director_profile'))
def global_semester_setup(request):
    if request.method == 'POST':
        academic_year = request.POST.get('academic_year')
        number = int(request.POST.get('number'))
        start_date = request.POST.get('start_date')
        end_date = request.POST.get('end_date')
        
        if not all([academic_year, number, start_date, end_date]):
            messages.error(request, "Заполните все поля!")
            return redirect('schedule:manage_semesters')

        try:
            with transaction.atomic():
                Semester.objects.filter(is_active=True).update(is_active=False)
                
                faculties = Faculty.objects.all()
                created_count = 0
                
                for faculty in faculties:
                    for course in range(1, 6):
                        shift = 'MORNING' if course <= 2 else 'DAY'
                        
                        Semester.objects.create(
                            faculty=faculty,
                            name=f"{'Осенний' if number == 1 else 'Весенний'} семестр",
                            academic_year=academic_year,
                            number=number,
                            course=course,
                            shift=shift,
                            start_date=start_date,
                            end_date=end_date,
                            is_active=True
                        )
                        created_count += 1
                
                messages.success(request, f"🚀 Успешно! Старые данные отправлены в архив. Создано {created_count} новых семестров. Система готова к работе с {start_date}!")
        except Exception as e:
            logger.exception("global_semester_setup")
            messages.error(request, _("Ошибка при генерации семестров. См. журнал сервера."))
            
        return redirect('schedule:manage_semesters')
        
    return render(request, 'schedule/global_setup.html')


@login_required
@user_passes_test(is_dean_or_admin)
def auto_schedule_config(request):
    semesters = Semester.objects.filter(is_active=True)
    institute = None
    
    if hasattr(request.user, 'dean_profile') or hasattr(request.user, 'vicedean_profile'):
        profile = getattr(request.user, 'dean_profile', None) or getattr(request.user, 'vicedean_profile', None)
        faculty = profile.faculty
        institute = faculty.institute if faculty else None
        
        semesters = semesters.filter(Q(faculty=faculty) | Q(faculty__isnull=True))
        groups = Group.objects.filter(specialty__department__faculty=faculty)
        
        if institute:
            teachers = Teacher.objects.filter(
                Q(department__faculty__institute=institute) | 
                Q(additional_departments__faculty__institute=institute) |
                Q(subject__department__faculty__institute=institute) |
                Q(subject__groups__specialty__department__faculty__institute=institute) |
                Q(scheduleslot__group__specialty__department__faculty__institute=institute)
            ).distinct()
            rooms = Classroom.objects.filter(building__institute=institute, is_active=True)
        else:
            teachers = Teacher.objects.filter(
                Q(department__faculty=faculty) | 
                Q(additional_departments__faculty=faculty) |
                Q(subject__department__faculty=faculty) |
                Q(subject__groups__specialty__department__faculty=faculty) |
                Q(scheduleslot__group__specialty__department__faculty=faculty)
            ).distinct()
            rooms = Classroom.objects.none()
            
    elif hasattr(request.user, 'director_profile') or hasattr(request.user, 'prorector_profile'):
        profile = getattr(request.user, 'director_profile', None) or getattr(request.user, 'prorector_profile', None)
        institute = profile.institute
        if institute:
            semesters = semesters.filter(Q(faculty__institute=institute) | Q(faculty__isnull=True))
            groups = Group.objects.filter(specialty__department__faculty__institute=institute)
            teachers = Teacher.objects.filter(
                Q(department__faculty__institute=institute) | 
                Q(additional_departments__faculty__institute=institute) |
                Q(subject__department__faculty__institute=institute) |
                Q(subject__groups__specialty__department__faculty__institute=institute) |
                Q(scheduleslot__group__specialty__department__faculty__institute=institute)
            ).distinct()
            rooms = Classroom.objects.filter(building__institute=institute, is_active=True)
        else:
            groups = Group.objects.all()
            teachers = Teacher.objects.all()
            rooms = Classroom.objects.filter(is_active=True)
    else:
        groups = Group.objects.all()
        teachers = Teacher.objects.all()
        rooms = Classroom.objects.filter(is_active=True)

    if request.method == 'POST':
        semester_id = request.POST.get('semester')
        group_ids = request.POST.getlist('groups')
        teacher_ids = request.POST.getlist('teachers')
        room_ids = request.POST.getlist('rooms')
        
        clear_existing = request.POST.get('clear_existing') == 'on'
        avoid_gaps = request.POST.get('avoid_gaps') == 'on'
        strict_room_types = request.POST.get('strict_room_types') == 'on'
        overflow_mode = int(request.POST.get('overflow_mode', 1))
        iterations = int(request.POST.get('iterations', 5))

        if not semester_id or not group_ids:
            messages.error(request, "Выберите семестр и хотя бы одну группу.")
            return redirect('schedule:auto_schedule_config')

        semester = get_object_or_404(Semester, id=semester_id)
        target_groups = Group.objects.filter(id__in=group_ids)
        
        if not institute and target_groups.exists():
            first_group = target_groups.first()
            if first_group.specialty and first_group.specialty.department.faculty:
                institute = first_group.specialty.department.faculty.institute
        
        try:
            with transaction.atomic():
                if clear_existing:
                    ScheduleSlot.objects.filter(
                        semester=semester,
                        group__in=target_groups,
                        is_military=False
                    ).delete()

                engine = AutoScheduleEngine(
                    semester=semester,
                    target_groups=target_groups,
                    target_teachers=teacher_ids if teacher_ids else None,
                    target_rooms=room_ids if room_ids else None,
                    avoid_gaps=avoid_gaps,
                    overflow_mode=overflow_mode,
                    strict_room_types=strict_room_types,
                    iterations=iterations,
                    institute=institute
                )
                result = engine.generate()

            if result['unassigned_count'] == 0:
                messages.success(request, f"✨ Успех! Сгенерировано {result['created']} занятий. Проанализировано {iterations} вариантов.")
            else:
                error_list = "<br>".join(result['unassigned_details'][:10])
                messages.warning(request, f"Сгенерировано {result['created']} занятий. Не удалось разместить {result['unassigned_count']} занятий (не хватило аудиторий или времени у преподавателей).<br><small>{error_list}</small>")
            
            if target_groups.exists():
                request.session['last_constructor_group'] = str(target_groups.first().id)
                request.session['last_constructor_semester'] = str(semester.id)
                return redirect(f"/schedule/constructor/?group={target_groups.first().id}&semester={semester.id}")
            return redirect('schedule:constructor')
            
        except Exception as e:
            logger.exception("auto_schedule_generate")
            messages.error(request, _("Критическая ошибка алгоритма. См. журнал сервера."))
            return redirect('schedule:auto_schedule_config')

    return render(request, 'schedule/auto_schedule.html', {
        'semesters': semesters,
        'groups': groups,
        'teachers': teachers,
        'rooms': rooms
    })



@login_required
@user_passes_test(is_dean_or_admin)
@login_required
@user_passes_test(is_dean_or_admin)
def manage_teacher_availability(request):
    teachers = Teacher.objects.select_related('user', 'department').all()
    
    if hasattr(request.user, 'dean_profile') or hasattr(request.user, 'vicedean_profile'):
        profile = getattr(request.user, 'dean_profile', None) or getattr(request.user, 'vicedean_profile', None)
        faculty = profile.faculty
        institute = faculty.institute if faculty else None
        if institute:
            teachers = teachers.filter(
                Q(department__faculty__institute=institute) | 
                Q(additional_departments__faculty__institute=institute) |
                Q(subject__department__faculty__institute=institute) |
                Q(subject__groups__specialty__department__faculty__institute=institute) |
                Q(scheduleslot__group__specialty__department__faculty__institute=institute)
            ).distinct()
        elif faculty:
            teachers = teachers.filter(
                Q(department__faculty=faculty) | 
                Q(additional_departments__faculty=faculty) |
                Q(subject__department__faculty=faculty) |
                Q(subject__groups__specialty__department__faculty=faculty) |
                Q(scheduleslot__group__specialty__department__faculty=faculty)
            ).distinct()
    elif hasattr(request.user, 'director_profile') or hasattr(request.user, 'prorector_profile'):
        profile = getattr(request.user, 'director_profile', None) or getattr(request.user, 'prorector_profile', None)
        institute = profile.institute
        if institute:
            teachers = teachers.filter(
                Q(department__faculty__institute=institute) | 
                Q(additional_departments__faculty__institute=institute) |
                Q(subject__department__faculty__institute=institute) |
                Q(subject__groups__specialty__department__faculty__institute=institute) |
                Q(scheduleslot__group__specialty__department__faculty__institute=institute)
            ).distinct()
        
    teacher_id = request.GET.get('teacher_id')
    selected_teacher = None
    unavailable_dict = {}
    
    time_slots = TimeSlot.objects.all().order_by('shift', 'start_time')
    days = [(0, 'Понедельник'), (1, 'Вторник'), (2, 'Среда'), (3, 'Четверг'), (4, 'Пятница'), (5, 'Суббота')]
    
    if teacher_id:
        selected_teacher = get_object_or_404(Teacher, id=teacher_id)
        if request.method == 'POST':
            TeacherUnavailableSlot.objects.filter(teacher=selected_teacher).delete()
            
            new_slots = []
            for key in request.POST.keys():
                if key.startswith('slot_'):
                    parts = key.split('_')
                    day = int(parts[1])
                    ts_id = int(parts[2])
                    new_slots.append(TeacherUnavailableSlot(
                        teacher=selected_teacher,
                        day_of_week=day,
                        time_slot_id=ts_id
                    ))
            
            if new_slots:
                TeacherUnavailableSlot.objects.bulk_create(new_slots)
                
            messages.success(request, f"График доступности для {selected_teacher.user.get_full_name()} успешно обновлен!")
            return redirect(f"{request.path}?teacher_id={teacher_id}")
            
        unavailable_qs = TeacherUnavailableSlot.objects.filter(teacher=selected_teacher)
        for u in unavailable_qs:
            if u.day_of_week not in unavailable_dict:
                unavailable_dict[u.day_of_week] = {}
            unavailable_dict[u.day_of_week][u.time_slot_id] = True

    return render(request, 'schedule/manage_teacher_availability.html', {
        'teachers': teachers,
        'selected_teacher': selected_teacher,
        'time_slots': time_slots,
        'days': days,
        'unavailable_dict': unavailable_dict
    })




@login_required
@require_POST
def api_create_credit_type(request):
    import json
    try:
        data = json.loads(request.body)
        name = data.get('name', '').strip()
        hours = int(data.get('hours', 24))
        faculty = request.user.dean_profile.faculty if hasattr(request.user, 'dean_profile') else None
        
        ct = CreditType.objects.create(name=name, hours_per_credit=hours, faculty=faculty)
        return JsonResponse({'success': True, 'id': ct.id, 'name': ct.name, 'hours': ct.hours_per_credit})
    except Exception as e:
        logger.exception("api_create_credit_type")
        return JsonResponse({'success': False, 'error': _('Внутренняя ошибка сервера')})

@login_required
@require_POST
def edit_plan_discipline(request, discipline_id):
    disc = get_object_or_404(PlanDiscipline, id=discipline_id)
    faculty = request.user.dean_profile.faculty if hasattr(request.user, 'dean_profile') else None
    form = PlanDisciplineForm(request.POST, instance=disc, faculty=faculty)
    if form.is_valid():
        form.save()
        messages.success(request, "Дисциплина успешно обновлена!")
    else:
        messages.error(request, "Ошибка при обновлении дисциплины.")
    return redirect(f"/schedule/plans/{disc.plan.id}/?semester={disc.semester_number}")


@user_passes_test(is_dean_or_admin)
def manage_credit_templates(request):
    templates = CreditTemplate.objects.all()
    if hasattr(request.user, 'dean_profile'):
        templates = templates.filter(Q(faculty=request.user.dean_profile.faculty) | Q(faculty__isnull=True))
    return render(request, 'schedule/manage_credit_templates.html', {'templates': templates})


@user_passes_test(is_dean_or_admin)
def add_credit_template(request):
    if request.method == 'POST':
        form = CreditTemplateForm(request.POST)
        if form.is_valid():
            template = form.save(commit=False)
            if hasattr(request.user, 'dean_profile'):
                template.faculty = request.user.dean_profile.faculty
            template.save()
            messages.success(request, _("Шаблон кредитов успешно добавлен!"))
            return redirect('schedule:manage_credit_templates')
    else:
        form = CreditTemplateForm()
    return render(request, 'core/form_generic.html', {
        'form': form, 'title': _('Добавить шаблон'),
        'cancel_url': reverse('schedule:manage_credit_templates')
    })


@user_passes_test(is_dean_or_admin)
def edit_credit_template(request, template_id):
    template = get_object_or_404(CreditTemplate, id=template_id)
    if request.method == 'POST':
        form = CreditTemplateForm(request.POST, instance=template)
        if form.is_valid():
            form.save()
            messages.success(request, _("Шаблон обновлен!"))
            return redirect('schedule:manage_credit_templates')
    else:
        form = CreditTemplateForm(instance=template)
    return render(request, 'core/form_generic.html', {
        'form': form, 'title': _('Редактировать шаблон'),
        'cancel_url': reverse('schedule:manage_credit_templates')
    })


@user_passes_test(is_dean_or_admin)
def delete_credit_template(request, template_id):
    template = get_object_or_404(CreditTemplate, id=template_id)
    template.delete()
    messages.success(request, _("Шаблон удален!"))
    return redirect('schedule:manage_credit_templates')


@login_required
@user_passes_test(is_dept_head_or_above)
@require_POST
def split_subject_load(request, subject_id):
    original_subject = get_object_or_404(Subject, id=subject_id)
    
    try:
        transfer_credits = int(request.POST.get('transfer_credits', 0))
        transfer_lec = int(request.POST.get('transfer_lec', 0))
        transfer_prac = int(request.POST.get('transfer_prac', 0))
        transfer_lab = int(request.POST.get('transfer_lab', 0))
        new_teacher_id = request.POST.get('new_teacher_id')
        
        if transfer_credits <= 0 and (transfer_lec + transfer_prac + transfer_lab) <= 0:
            messages.error(request, _("Укажите часы или кредиты для переноса."))
            return redirect('schedule:manage_subjects')

        with transaction.atomic():
            original_subject.credits = max(0, original_subject.credits - transfer_credits)
            original_subject.lecture_hours = max(0, original_subject.lecture_hours - transfer_lec)
            original_subject.practice_hours = max(0, original_subject.practice_hours - transfer_prac)
            original_subject.lab_hours = max(0, original_subject.lab_hours - transfer_lab)
            original_subject.save()

            new_teacher = Teacher.objects.filter(id=new_teacher_id).first() if new_teacher_id else None
            
            new_subject = Subject.objects.create(
                name=f"{original_subject.name} (Часть 2)",
                code=f"{original_subject.code}_split_{uuid.uuid4().hex[:4]}",
                department=original_subject.department,
                type=original_subject.type,
                credits=transfer_credits,
                lecture_hours=transfer_lec,
                practice_hours=transfer_prac,
                lab_hours=transfer_lab,
                control_hours=0,
                independent_work_hours=0,
                semester_weeks=original_subject.semester_weeks,
                teacher=new_teacher,
                plan_discipline=original_subject.plan_discipline,
                is_active=True
            )
            new_subject.groups.set(original_subject.groups.all())
            
        messages.success(request, _("Нагрузка успешно разделена. Создана новая карточка предмета."))
    except Exception as e:
        logger.exception("split_subject_load")
        messages.error(request, _("Ошибка при разделении нагрузки. См. журнал сервера."))
        
    return redirect('schedule:manage_subjects')


@login_required
@user_passes_test(is_dept_head_or_above)
def toggle_subject_active(request, subject_id):
    subject = get_object_or_404(Subject, id=subject_id)
    subject.is_active = not getattr(subject, 'is_active', True)
    subject.save()
    status = _("активирована") if subject.is_active else _("деактивирована (исключена из нагрузки)")
    messages.success(request, _("Дисциплина ") + subject.name + " " + status + ".")
    return redirect('schedule:manage_subjects')


@login_required
@user_passes_test(lambda u: u.is_superuser or u.role in [
    'HEAD_OF_DEPT', 'DEAN', 'VICE_DEAN', 'DIRECTOR', 'PRO_RECTOR'
])
def import_department_load(request):
    def _get_dept_for_user():
        user = request.user
        if hasattr(user, 'head_of_dept_profile') and user.head_of_dept_profile.department:
            return user.head_of_dept_profile.department
        if hasattr(user, 'dean_profile') and user.dean_profile.faculty:
            return Department.objects.filter(
                faculty=user.dean_profile.faculty
            ).first()
        if hasattr(user, 'vicedean_profile') and user.vicedean_profile.faculty:
            return Department.objects.filter(
                faculty=user.vicedean_profile.faculty
            ).first()
        return Department.objects.first()

    def _get_teachers_for_user():
        user = request.user
        if hasattr(user, 'head_of_dept_profile') and user.head_of_dept_profile.department:
            return Teacher.objects.filter(
                department=user.head_of_dept_profile.department
            ).select_related('user')
        if hasattr(user, 'dean_profile') and user.dean_profile.faculty:
            return Teacher.objects.filter(
                department__faculty=user.dean_profile.faculty
            ).select_related('user')
        if hasattr(user, 'vicedean_profile') and user.vicedean_profile.faculty:
            return Teacher.objects.filter(
                department__faculty=user.vicedean_profile.faculty
            ).select_related('user')
        return Teacher.objects.select_related('user').all()

    if request.method == 'POST' and 'confirm_import' in request.POST:
        preview_data = request.session.get('dept_load_preview', [])
        if not preview_data:
            messages.error(request, "Сессия истекла. Загрузите файл заново.")
            return redirect('schedule:import_department_load')

        dept = _get_dept_for_user()
        created_count = 0
        errors = []

        try:
            with transaction.atomic():
                for item in preview_data:
                    idx = str(item['id'])
                    group_id = item.get('group_id')
                    if not group_id:
                        continue

                    subj_name = request.POST.get(f'subject_{idx}', item.get('subject', '')).strip()
                    if not subj_name:
                        continue

                    teacher_id_raw = request.POST.get(f'teacher_{idx}')
                    teacher_obj = None
                    if teacher_id_raw and str(teacher_id_raw).isdigit():
                        teacher_obj = Teacher.objects.filter(id=int(teacher_id_raw)).first()

                    credits = safe_int(request.POST.get(f'credits_{idx}', item.get('credits', 0)))
                    lec     = safe_int(request.POST.get(f'lec_{idx}',     item.get('lec', 0)))
                    prac    = safe_int(request.POST.get(f'prac_{idx}',    item.get('prac', 0)))
                    lab     = safe_int(request.POST.get(f'lab_{idx}',     item.get('lab', 0)))

                    try:
                        group_obj = Group.objects.get(id=group_id)
                    except Group.DoesNotExist:
                        errors.append(f"Группа id={group_id} не найдена, строка пропущена.")
                        continue

                    subject = Subject.objects.create(
                        name=subj_name,
                        code=f"IMP_{uuid.uuid4().hex[:8]}",
                        department=dept,
                        credits=credits,
                        lecture_hours=lec,
                        practice_hours=prac,
                        lab_hours=lab,
                        teacher=teacher_obj,
                        is_active=True
                    )
                    subject.groups.add(group_obj)
                    created_count += 1

        except Exception as e:
            messages.error(request, f"Критическая ошибка при сохранении: {e}")
            return redirect('schedule:import_department_load')
        finally:
            request.session.pop('dept_load_preview', None)

        if errors:
            for err in errors:
                messages.warning(request, err)
        if created_count > 0:
            messages.success(
                request,
                f"Успешно импортировано и распределено {created_count} дисциплин."
            )
        else:
            messages.warning(request, "Ни одна дисциплина не была создана. Проверьте данные.")
        return redirect('schedule:manage_subjects')

    if request.method == 'POST' and 'excel_file' in request.FILES:
        excel_file = request.FILES['excel_file']
        preview_data = []
        parse_errors = []

        try:
            import openpyxl
            wb = openpyxl.load_workbook(excel_file, data_only=True)
            sheet = wb.active

            for row_idx, row in enumerate(
                sheet.iter_rows(min_row=2, values_only=True), start=2
            ):
                if not row or all(c is None for c in row):
                    continue

                subj_name = str(row[0] or '').strip()
                if not subj_name or subj_name.lower() in ('none', ''):
                    continue

                group_name   = str(row[1]).strip() if len(row) > 1 and row[1] else ""
                teacher_name = str(row[6]).strip() if len(row) > 6 and row[6] else ""

                group_obj   = Group.objects.filter(name__iexact=group_name).first()
                teacher_obj = None
                if teacher_name:
                    first_word = teacher_name.split()[0]
                    teacher_obj = Teacher.objects.filter(
                        user__last_name__icontains=first_word
                    ).first()

                if not group_obj and group_name:
                    parse_errors.append(
                        f"Строка {row_idx}: группа «{group_name}» не найдена в базе."
                    )

                preview_data.append({
                    'id':          row_idx,
                    'subject':     subj_name,
                    'group':       group_name,
                    'group_id':    group_obj.id if group_obj else None,
                    'credits':     safe_int(row[2]) if len(row) > 2 else 0,
                    'lec':         safe_int(row[3]) if len(row) > 3 else 0,
                    'prac':        safe_int(row[4]) if len(row) > 4 else 0,
                    'lab':         safe_int(row[5]) if len(row) > 5 else 0,
                    'teacher_name': teacher_name,
                    'teacher_id':  teacher_obj.id if teacher_obj else None,
                })

        except Exception as e:
            messages.error(request, f"Ошибка чтения файла: {e}")
            return redirect('schedule:import_department_load')

        if not preview_data:
            messages.warning(request, "Файл не содержит данных (начиная со 2-й строки).")
            return redirect('schedule:import_department_load')

        request.session['dept_load_preview'] = preview_data
        request.session.modified = True

        for err in parse_errors:
            messages.warning(request, err)

        teachers = _get_teachers_for_user()
        return render(request, 'schedule/import_dept_load_preview.html', {
            'preview_data': preview_data,
            'teachers': teachers,
            'parse_warnings_count': len(parse_errors),
        })

    return render(request, 'schedule/import_dept_load.html')

@login_required
@require_POST
def clear_military_day(request):
    if not is_dean_or_admin(request.user):
        return JsonResponse({'success': False, 'error': 'Нет прав'}, status=403)
    try:
        data = json.loads(request.body)
        group_id   = data.get('group_id')
        semester_id = data.get('semester_id')
        day_of_week = data.get('day_of_week')
 
        group    = get_object_or_404(Group, id=group_id)
        semester = get_object_or_404(Semester, id=semester_id)
 
        if hasattr(request.user, 'dean_profile'):
            faculty = request.user.dean_profile.faculty
            if group.specialty and group.specialty.department.faculty != faculty:
                return JsonResponse({'success': False, 'error': 'Нет доступа к этой группе'}, status=403)
 
        deleted, _ = ScheduleSlot.objects.filter(
            group=group,
            semester=semester,
            day_of_week=day_of_week,
            is_military=True,
        ).delete()
 
        return JsonResponse({'success': True, 'deleted': deleted})
    except Exception as e:
        logger.exception("clear_military_day")
        return JsonResponse({'success': False, 'error': _('Внутренняя ошибка сервера')}, status=500)



@login_required
def schedule_calendar(request):
    user = request.user

    if user.is_superuser:
        groups = list(Group.objects.all().order_by('course', 'name'))
        selected_group = None
        
    elif user.role == 'STUDENT' and hasattr(user, 'student_profile'):
        group = user.student_profile.group if user.student_profile else None
        groups = [group] if group else []
        selected_group = group
        
    elif user.role == 'TEACHER' and hasattr(user, 'teacher_profile'):
        group_ids = ScheduleSlot.objects.filter(
            teacher=user.teacher_profile, is_active=True
        ).values_list('group_id', flat=True).distinct()
        groups = list(Group.objects.filter(id__in=group_ids).order_by('course', 'name'))
        selected_group = groups[0] if groups else None
        
    elif is_dean_or_admin(user):
        if hasattr(user, 'dean_profile') and getattr(user.dean_profile, 'faculty', None):
            groups = list(Group.objects.filter(
                specialty__department__faculty=user.dean_profile.faculty
            ).order_by('course', 'name'))
        elif hasattr(user, 'director_profile') or hasattr(user, 'prorector_profile'):
            profile = getattr(user, 'director_profile', None) or getattr(user, 'prorector_profile', None)
            institute = getattr(profile, 'institute', None)
            if institute:
                groups = list(Group.objects.filter(
                    specialty__department__faculty__institute=institute
                ).order_by('course', 'name'))
            else:
                groups = list(Group.objects.all().order_by('course', 'name'))
        else:
            groups = list(Group.objects.all().order_by('course', 'name'))
        selected_group = None
        
    else:
        groups = []
        selected_group = None

    group_id = request.GET.get('group')
    if group_id:
        selected_group = Group.objects.filter(id=group_id).first()

    active_semester = None
    if selected_group:
        active_semester = get_active_semester_for_group(selected_group)
    if not active_semester:
        active_semester = Semester.objects.filter(is_active=True).first()

    can_edit = is_dean_or_admin(user)
    can_move_calendar = (
        is_dean_or_admin(user)
        or getattr(user, 'teacher_profile', None)
        or getattr(user, 'head_of_dept_profile', None)
    )

    weekly_slots_count = 0
    unused_hours = []
    if selected_group and active_semester:
        slots = ScheduleSlot.objects.filter(
            group=selected_group, semester=active_semester, is_active=True
        )
        weekly_slots_count = sum([1 if s.week_type == 'EVERY' else 0.5 for s in slots])

        unused_hours = UnusedHourPool.objects.filter(
            group=selected_group,
            is_recovered=False
        ).select_related('subject', 'teacher__user').order_by('-original_date')

    return render(request, 'schedule/schedule_calendar.html', {
        'groups': groups,
        'selected_group': selected_group,
        'active_semester': active_semester,
        'can_edit': can_edit,
        'can_move_calendar': can_move_calendar,
        'calendar_user_id': user.id,
        'calendar_is_teacher': hasattr(user, 'teacher_profile'),
        'is_teacher_view': hasattr(user, 'teacher_profile') and not request.GET.get('group'),
        'weekly_slots_count': weekly_slots_count,
        'unused_hours': unused_hours,
    })



 
 
@login_required
def calendar_sidebar_api(request):
    try:
        group_id = request.GET.get('group')
        logger.info(
            f"--- ЗАПРОС ОБНОВЛЕНИЯ БОКОВОЙ ПАНЕЛИ: Пользователь {request.user.username}, группа {group_id} ---"
        )

        if not group_id:
            return JsonResponse({'weekly_slots': 0, 'unused_hours': []})

        group = Group.objects.filter(id=group_id).first()
        if not group:
            return JsonResponse({'weekly_slots': 0, 'unused_hours': []})

        active_semester = get_active_semester_for_group(group)
        if not active_semester:
            return JsonResponse({'weekly_slots': 0, 'unused_hours': []})

        slots = ScheduleSlot.objects.filter(group=group, semester=active_semester, is_active=True)
        weekly_slots_count = sum([1 if s.week_type == 'EVERY' else 0.5 for s in slots])

        unused = UnusedHourPool.objects.filter(
            group=group,
            is_recovered=False
        ).select_related('subject', 'teacher__user').order_by('-original_date')

        unused_data = [{
            'id': u.id,
            'subject': u.subject.name,
            'teacher': u.teacher.user.get_full_name() if u.teacher and u.teacher.user else '-',
            'date': u.original_date.strftime("%d.%m.%Y") if u.original_date else '-'
        } for u in unused]

        logger.info(
            f"Боковая панель успешно собрана: недельных пар {weekly_slots_count}, банк отмен: {len(unused_data)}"
        )
        return JsonResponse({'weekly_slots': weekly_slots_count, 'unused_hours': unused_data})
    except Exception as e:
        logger.exception(f"КРИТИЧЕСКАЯ ОШИБКА В API БОКОВОЙ ПАНЕЛИ: {str(e)}")
        return JsonResponse({'weekly_slots': 0, 'unused_hours': []}, status=500)


@login_required
def schedule_calendar_events(request):
    group_id = request.GET.get('group')
    start_str = request.GET.get('start', '')[:10]
    end_str   = request.GET.get('end',   '')[:10]
 
    try:
        start_date = date_cls.fromisoformat(start_str)
        end_date   = date_cls.fromisoformat(end_str)
    except Exception:
        start_date = date_cls.today()
        end_date   = start_date + timedelta(days=35)
 
    user = request.user
 
    if group_id:
        slots_qs = ScheduleSlot.objects.filter(group_id=group_id, is_active=True)
    elif hasattr(user, 'student_profile') and getattr(user.student_profile, 'group', None):
        slots_qs = ScheduleSlot.objects.filter(group=user.student_profile.group, is_active=True)
    elif hasattr(user, 'teacher_profile'):
        slots_qs = ScheduleSlot.objects.filter(teacher=user.teacher_profile, is_active=True)
    else:
        slots_qs = ScheduleSlot.objects.none()
 
    slots_qs = slots_qs.select_related('subject', 'teacher__user', 'group', 'time_slot', 'semester', 'classroom')
 
    slot_ids = list(slots_qs.values_list('id', flat=True))
    exceptions_qs = SchedExc.objects.filter(
        schedule_slot_id__in=slot_ids,
        exception_date__gte=start_date,
        exception_date__lte=end_date + timedelta(days=7),
    )
    exc_map = {(e.schedule_slot_id, e.exception_date): e for e in exceptions_qs}
 
    COLORS = {'LECTURE': '#2563eb', 'PRACTICE': '#16a34a', 'LAB': '#0891b2', 'SRSP': '#d97706'}
    events =[]
 
    for slot in slots_qs:
        sem = slot.semester
        if not sem or not sem.start_date or not sem.end_date: continue
 
        vac_weeks =[int(w.strip()) for w in sem.vacation_weeks.split(',') if w.strip().isdigit()] if sem.vacation_weeks else[]
        loop_start = max(sem.start_date, start_date)
        days_ahead  = (slot.day_of_week - loop_start.weekday()) % 7
        current     = loop_start + timedelta(days=days_ahead)
 
        while current <= min(sem.end_date, end_date):
            week_num = ((current - sem.start_date).days // 7) + 1
 
            if week_num in vac_weeks:
                current += timedelta(weeks=1)
                continue
 
            is_red = (week_num % 2 == 1)
            show = (slot.week_type == 'EVERY' or (slot.week_type == 'RED' and is_red) or (slot.week_type == 'BLUE' and not is_red))
            
            if not show:
                current += timedelta(weeks=1)
                continue
 
            exc = exc_map.get((slot.id, current))
 
            if exc and exc.exception_type == 'CANCEL':
                events.append({
                    'id': f'exc_cancel_{exc.id}',
                    'title': f'❌ {slot.subject.name}',
                    'start': f'{current}T{slot.time_slot.start_time.strftime("%H:%M:%S")}',
                    'end': f'{current}T{slot.time_slot.end_time.strftime("%H:%M:%S")}',
                    'color': '#9ca3af',
                    'classNames':['fc-cancelled'],
                    'extendedProps': {
                        'cancelled': True, 'reason': exc.reason, 'exception_id': exc.id,
                        'type_display': slot.get_lesson_type_display(),
                        'group': slot.group.name,
                        'teacher': slot.teacher.user.get_full_name() if slot.teacher else '—',
                        'teacher_id': slot.teacher.user_id if slot.teacher else None,
                        'room': slot.room or '—', 'slot_id': slot.id,
                    }
                })
                current += timedelta(weeks=1)
                continue
 
            if exc and exc.exception_type == 'RESCHEDULE' and exc.new_date:
                new_st = exc.new_start_time or slot.time_slot.start_time
                new_et = exc.new_end_time or slot.time_slot.end_time
                if start_date <= exc.new_date <= end_date:
                    events.append({
                        'id': f'exc_moved_{exc.id}',
                        'title': f'🔄 {slot.subject.name}',
                        'start': f'{exc.new_date}T{new_st.strftime("%H:%M:%S")}',
                        'end': f'{exc.new_date}T{new_et.strftime("%H:%M:%S")}',
                        'color': COLORS.get(slot.lesson_type, '#6b7280'),
                        'editable': True,
                        'extendedProps': {
                            'rescheduled': True, 'reason': exc.reason, 'exception_id': exc.id,
                            'original_date': str(current),
                            'type_display': slot.get_lesson_type_display(),
                            'group': slot.group.name,
                            'teacher': slot.teacher.user.get_full_name() if slot.teacher else '—',
                            'teacher_id': slot.teacher.user_id if slot.teacher else None,
                            'room': slot.room or '—', 'slot_id': slot.id,
                        }
                    })
                current += timedelta(weeks=1)
                continue
 
            events.append({
                'id': f'slot_{slot.id}_{current}',
                'title': ('🪖 Воен. каф' if slot.is_military else slot.subject.name),
                'start': f'{current}T{slot.time_slot.start_time.strftime("%H:%M:%S")}',
                'end': f'{current}T{slot.time_slot.end_time.strftime("%H:%M:%S")}',
                'backgroundColor': '#1f2937' if slot.is_military else COLORS.get(slot.lesson_type, '#6b7280'),
                'borderColor': '#1f2937' if slot.is_military else COLORS.get(slot.lesson_type, '#6b7280'),
                'textColor': '#ffffff',
                'editable': True,
                'extendedProps': {
                    'slot_id': slot.id, 'original_date': str(current),
                    'type_display': slot.get_lesson_type_display(),
                    'group': slot.group.name, 'teacher': slot.teacher.user.get_full_name() if slot.teacher else '—',
                    'teacher_id': slot.teacher.user_id if slot.teacher else None,
                    'room': slot.room or (slot.classroom.number if slot.classroom else '—'),
                    'is_stream': bool(slot.stream_id), 'military': slot.is_military,
                    'lesson_type': slot.lesson_type,
                }
            })
            current += timedelta(weeks=1)

    if hasattr(user, 'student_profile') or hasattr(user, 'teacher_profile'):
        assignments = Assignment.objects.filter(
            due_date__gte=start_date, due_date__lte=end_date,
            module__section__course__enrolments__user=user
        ).select_related('module__section__course')
        
        for ass in assignments:
            events.append({
                'id': f'lms_{ass.id}',
                'title': f'📚 LMS Дедлайн: {ass.module.title} ({ass.module.section.course.short_name})',
                'start': ass.due_date.isoformat(),
                'color': '#8b5cf6', 
                'extendedProps': {
                    'is_lms': True, 'url': f'/lms/modules/{ass.module.id}/'
                }
            })

    return JsonResponse(events, safe=False)
 
 
@login_required
@require_POST
def calendar_move_slot(request):
    import json
    logger.info(f"--- ПЕРЕНОС/ОТМЕНА СЛОТА: Пользователь {request.user.username} ---")
    
    try:
        data = json.loads(request.body)
        logger.info(f"Данные запроса: {data}")
        move_type = data.get('move_type', 'once')

        if move_type == 'recover':
            pool_id = data.get('pool_id')
            new_date = date_cls.fromisoformat(str(data['new_date'])[:10])
            new_start = dt_cls.strptime(data['new_start_time'][:5], '%H:%M').time()

            logger.info(
                f"Попытка отработки занятия. Pool ID: {pool_id}, Новая дата: {new_date}, Время: {new_start}"
            )

            pool_item = get_object_or_404(UnusedHourPool, id=pool_id)

            with transaction.atomic():
                exc = SchedExc.objects.filter(
                    schedule_slot__group=pool_item.group,
                    schedule_slot__subject=pool_item.subject,
                    exception_date=pool_item.original_date,
                    exception_type='CANCEL'
                ).first()

                if exc:
                    logger.info(f"Найдено исключение отмены ID: {exc.id}. Конвертируем в RESCHEDULE.")
                    exc.exception_type = 'RESCHEDULE'
                    exc.new_date = new_date
                    exc.new_start_time = new_start

                    target_ts = TimeSlot.objects.filter(start_time__lte=new_start).order_by('-start_time').first()
                    if target_ts:
                        exc.new_end_time = target_ts.end_time
                    else:
                        exc.new_end_time = (datetime.combine(date_cls.today(), new_start) + timedelta(minutes=50)).time()

                    exc.save()
                    pool_item.is_recovered = True
                    pool_item.save()
                    logger.info("Занятие успешно восстановлено из банка часов.")
                    return JsonResponse({'success': True, 'type': 'recovered'})
                else:
                    logger.warning("Оригинальное исключение CANCEL не найдено. Базовый слот мог быть удален полностью.")
                    return JsonResponse({
                        'success': False,
                        'error': 'Невозможно восстановить: исходное занятие было удалено из общей сетки расписания (а не просто отменено на 1 день).'
                    })

        if move_type == 'revert':
            exc_id = data.get('exception_id')
            logger.info(f"Откат исключения ID: {exc_id}")
            exc = SchedExc.objects.filter(id=exc_id).first()
            if exc:
                with transaction.atomic():
                    UnusedHourPool.objects.filter(
                        group=exc.schedule_slot.group,
                        subject=exc.schedule_slot.subject,
                        original_date=exc.exception_date
                    ).delete()
                    exc.delete()
            return JsonResponse({'success': True, 'type': 'reverted'})

        try:
            slot_id = int(data.get('slot_id'))
            slot = get_object_or_404(ScheduleSlot, id=slot_id)
        except (ValueError, TypeError):
            logger.error(f"Некорректный slot_id: {data.get('slot_id')}")
            return JsonResponse({'success': False, 'error': 'Некорректный ID занятия. Дождитесь загрузки календаря.'})
        user = request.user

        safe_date_str = str(data['original_date'])[:10] 
        original_date = date_cls.fromisoformat(safe_date_str)
        reason = data.get('reason', 'Изменение в живом календаре')

        if move_type == 'cancel':
            logger.info(f"Одиночная отмена слота ID: {slot.id} на дату {original_date}")
            with transaction.atomic():
                SchedExc.objects.update_or_create(
                    schedule_slot=slot, exception_date=original_date,
                    defaults={'exception_type': 'CANCEL', 'reason': reason}
                )
                UnusedHourPool.objects.get_or_create(
                    group=slot.group, subject=slot.subject, teacher=slot.teacher,
                    semester=slot.semester, original_date=original_date,
                    defaults={'reason': reason}
                )
            logger.info("Отмена успешно сохранена в Банк Часов")
            return JsonResponse({'success': True, 'type': 'cancelled'})

        new_date = date_cls.fromisoformat(str(data['new_date'])[:10])
        new_start = dt_cls.strptime(data['new_start_time'][:5], '%H:%M').time()
        new_end = dt_cls.strptime(data['new_end_time'][:5], '%H:%M').time()
        
        logger.info(f"Перенос на {new_date} с {new_start} до {new_end}")

        target_ts = TimeSlot.objects.filter(start_time__lte=new_start).order_by('-start_time').first()
        if not target_ts:
            target_ts = TimeSlot.objects.first()

        if move_type == 'once':
            SchedExc.objects.update_or_create(
                schedule_slot=slot, exception_date=original_date,
                defaults={
                    'exception_type': 'RESCHEDULE', 'reason': reason,
                    'new_date': new_date, 'new_start_time': new_start,
                    'new_end_time': new_end, 'new_classroom': slot.classroom
                }
            )
            return JsonResponse({'success': True, 'type': 'exception_created'})
     
        elif move_type == 'all':
            update_fields = dict(day_of_week=new_date.weekday(), time_slot=target_ts,
                                 start_time=new_start, end_time=new_end)
            ScheduleSlot.objects.filter(id=slot.id).update(**update_fields)
            if slot.stream_id:
                ScheduleSlot.objects.filter(stream_id=slot.stream_id).update(**update_fields)
            return JsonResponse({'success': True, 'type': 'slot_updated'})
            
    except Exception as e:
        logger.exception(f"ОШИБКА ПРИ ПЕРЕНОСЕ/ОТМЕНЕ: {str(e)}")
        return JsonResponse({'success': False, 'error': str(e)})





@login_required
@require_POST
def mass_cancel_day(request):
    import json
    logger.info(f"=== МАССОВАЯ ОТМЕНА ЗАПУЩЕНА ПОЛЬЗОВАТЕЛЕМ: {request.user.username} +++")
    
    try:
        data = json.loads(request.body)
        logger.info(f"Получены данные: {data}")
        
        cancel_date_str = data.get('date')
        reason = data.get('reason', 'Массовая отмена занятий')

        if not cancel_date_str:
            logger.warning("Дата не указана в запросе")
            return JsonResponse({'success': False, 'error': 'Дата не указана'})

        cancel_date = dt_cls.fromisoformat(str(cancel_date_str)[:10]).date()
        day_of_week = cancel_date.weekday()
        user = request.user

        logger.info(f"Дата отмены: {cancel_date}, День недели: {day_of_week}")

        slots = ScheduleSlot.objects.filter(day_of_week=day_of_week, is_active=True)

        if user.is_superuser or hasattr(user, 'director_profile') or hasattr(user, 'prorector_profile'):
            logger.info("Права: Глобальный админ")
        elif hasattr(user, 'dean_profile') and user.dean_profile.faculty:
            logger.info(f"Права: Декан факультета {user.dean_profile.faculty.name}")
            slots = slots.filter(group__specialty__department__faculty=user.dean_profile.faculty)
        elif hasattr(user, 'head_of_dept_profile') and user.head_of_dept_profile.department:
            logger.info(f"Права: Зав. кафедрой {user.head_of_dept_profile.department.name}")
            slots = slots.filter(subject__department=user.head_of_dept_profile.department)
        else:
            logger.warning("Отказ в доступе: Нет прав")
            return JsonResponse({'success': False, 'error': 'Нет прав на массовую отмену'})

        slots_count = slots.count()
        logger.info(f"Найдено слотов для отмены: {slots_count}")

        with transaction.atomic():
            count = 0
            for slot in slots:
                sem = slot.semester
                if not sem or not sem.start_date: continue
                
                week_num = ((cancel_date - sem.start_date).days // 7) + 1
                is_red = (week_num % 2 == 1)
                show = (slot.week_type == 'EVERY' or (slot.week_type == 'RED' and is_red) or (slot.week_type == 'BLUE' and not is_red))

                if show:
                    logger.info(f"Отменяем слот ID: {slot.id}, Предмет: {slot.subject.name}")
                    SchedExc.objects.update_or_create(
                        schedule_slot=slot, exception_date=cancel_date,
                        defaults={'exception_type': 'CANCEL', 'reason': reason}
                    )
                    UnusedHourPool.objects.get_or_create(
                        group=slot.group, subject=slot.subject, teacher=slot.teacher,
                        semester=slot.semester, original_date=cancel_date,
                        defaults={'reason': reason}
                    )
                    count += 1

        logger.info(f"=== УСПЕШНО ОТМЕНЕНО: {count} ПАР ===")
        return JsonResponse({'success': True, 'count': count})
        
    except Exception as e:
        logger.exception(f"КРИТИЧЕСКАЯ ОШИБКА В MASS_CANCEL: {str(e)}")
        return JsonResponse({'success': False, 'error': f"Внутренняя ошибка: {str(e)}"})

